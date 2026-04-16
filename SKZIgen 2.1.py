#!/usr/bin/env python3
"""
Генератор заявок ViPNet PKI Client
Установка: pip3 install python-docx openpyxl xlrd
Запуск:    python3 pki_zayavka.py
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json, os, subprocess, sys, copy, zipfile, re, socket
from datetime import datetime
import xlrd
import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".pki_zayavka_config.json")

MONTHS_RU = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}

SMALL_WORDS = {
    "и",
    "или",
    "а",
    "но",
    "для",
    "с",
    "на",
    "по",
    "в",
    "из",
    "до",
    "при",
    "за",
    "над",
    "под",
    "об",
    "о",
    "к",
    "у",
    "от",
    "то",
    "же",
    "бы",
}

PURPOSE_OPTIONS = ["для подписания ЭП на портале ЕЦП"]
PURPOSE_CUSTOM = "— ввести свой вариант —"

PC_HEADERS = [
    "этаж",
    "№ кабинета",
    "Имя ПК",
    "доменное имя",
    "ip-адрес",
    "ФИО пользователя",
    "срок действия серт.",
    "СКЗИ",
    "Серийный номер на корпусе",
    "Инвентарный номер",
    "№ стикера",
    "",
]
NUM_PC_COLS = len(PC_HEADERS)  # 12


# ══════════════════════════════════════════════════════
#  Тема Excel → реальные цвета
# ══════════════════════════════════════════════════════

# Стандартная схема цветов Office (из theme1.xml файла актуализации)
# Порядок: dk1, lt1, dk2, lt2, accent1-6, hlink, folHlink
THEME_COLORS = [
    "#000000",  # 0  dk1  (windowText)
    "#FFFFFF",  # 1  lt1  (window)
    "#1F497D",  # 2  dk2
    "#EEECE1",  # 3  lt2
    "#4F81BD",  # 4  accent1
    "#C0504D",  # 5  accent2
    "#9BBB59",  # 6  accent3
    "#8064A2",  # 7  accent4
    "#4BACC6",  # 8  accent5
    "#F79646",  # 9  accent6  (оранжевый, с tint → коричневый)
    "#0000FF",  # 10 hlink
    "#800080",  # 11 folHlink
]


def _load_theme_from_xlsx(path: str) -> list:
    """Пытается извлечь цвета темы прямо из xlsx. Возвращает THEME_COLORS если не получилось."""
    try:
        with zipfile.ZipFile(path) as z:
            theme_files = [
                f for f in z.namelist() if re.search(r"theme/theme\d*\.xml$", f)
            ]
            if not theme_files:
                return THEME_COLORS
            content = z.read(theme_files[0]).decode("utf-8")
        # Порядок тегов в clrScheme: dk1, lt1, dk2, lt2, accent1..6, hlink, folHlink
        colors = []
        for m in re.finditer(
            r"<a:(?:dk1|lt1|dk2|lt2|accent\d|hlink|folHlink)>"
            r'.*?(?:<a:srgbClr val="([0-9A-Fa-f]{6})"'
            r'|<a:sysClr[^>]+lastClr="([0-9A-Fa-f]{6})")',
            content,
        ):
            rgb = m.group(1) or m.group(2)
            colors.append("#" + rgb.upper())
        return colors if len(colors) >= 10 else THEME_COLORS
    except Exception:
        return THEME_COLORS


def _apply_tint(hex_color: str, tint: float) -> str:
    """Применяет tint к hex-цвету: tint<0 темнее, tint>0 светлее."""
    if not hex_color or tint == 0.0:
        return hex_color
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    if tint < 0:
        r = int(r * (1 + tint))
        g = int(g * (1 + tint))
        b = int(b * (1 + tint))
    else:
        r = int(r + (255 - r) * tint)
        g = int(g + (255 - g) * tint)
        b = int(b + (255 - b) * tint)
    return f"#{max(0,min(255,r)):02X}{max(0,min(255,g)):02X}{max(0,min(255,b)):02X}"


def _resolve_color(color_obj, theme_palette: list) -> str:
    """Превращает openpyxl Color в '#RRGGBB' или '' если нет цвета."""
    if color_obj is None:
        return ""
    try:
        if color_obj.type == "rgb":
            argb = color_obj.rgb  # 'FFRRGGBB' или '00000000'
            if argb in ("00000000", ""):
                return ""
            return "#" + argb[2:].upper()
        elif color_obj.type == "theme":
            idx = color_obj.theme
            tint = color_obj.tint or 0.0
            if idx < len(theme_palette):
                return _apply_tint(theme_palette[idx], tint)
        elif color_obj.type == "indexed":
            # Стандартная indexed палитра (первые 8 — основные)
            INDEXED = [
                "#000000",
                "#FFFFFF",
                "#FF0000",
                "#00FF00",
                "#0000FF",
                "#FFFF00",
                "#FF00FF",
                "#00FFFF",
            ]
            idx = color_obj.indexed
            if idx < len(INDEXED):
                return INDEXED[idx]
    except Exception:
        pass
    return ""


def _readable_fg(bg: str, fg: str) -> str:
    """Возвращает читаемый цвет текста. Не меняет цвет если фон тёмный."""
    if not fg:
        return "#000000"
    # Определяем яркость фона (пустой bg = белый/светлый)
    if bg:
        r = int(bg[1:3], 16)
        g = int(bg[3:5], 16)
        b = int(bg[5:7], 16)
        luminance = 0.299 * r + 0.587 * g + 0.114 * b
        bg_is_dark = luminance < 140
    else:
        bg_is_dark = False  # нет заливки = светлый фон

    # Белый и жёлтый текст нечитаем на светлом фоне
    if not bg_is_dark and fg.upper() in ("#FFFFFF", "#FFFF00"):
        return "#000000"
    return fg


# ══════════════════════════════════════════════════════
#  Конфиг
# ══════════════════════════════════════════════════════


def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_config(cfg):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════
#  Форматирование
# ══════════════════════════════════════════════════════


def make_initials(full_name: str) -> str:
    parts = full_name.strip().split()
    if len(parts) >= 2:
        return parts[0] + " " + "".join(p[0] + "." for p in parts[1:])
    return full_name


def abbreviate_dept(name: str) -> str:
    if not name:
        return name
    parts = name.split()
    result = []
    i = 0
    while i < len(parts):
        w = parts[i]
        if w == "№" and i + 1 < len(parts):
            result.append(" №" + parts[i + 1])
            i += 2
            continue
        if w.isdigit():
            result.append(w)
        elif "-" not in w and w.lower() in SMALL_WORDS:
            result.append(w.lower())
        elif w[0].isalpha():
            for sp in w.split("-"):
                if sp and sp[0].isalpha():
                    result.append(sp[0].upper())
        i += 1
    return "".join(result)


def build_position_doc(position: str, department: str) -> str:
    if not department:
        return position
    pw = position.split()
    dw = department.split()
    if pw and dw:
        if pw[-1].lower()[:5] == dw[0].lower()[:5]:
            return position + " " + " ".join(dw[1:])
        else:
            return position + " " + department[0].lower() + department[1:]
    return position + " " + department


def chief_position_prefix(chief_position: str) -> str:
    if not chief_position:
        return ""
    dept_kw = {"отдела", "управления", "службы", "сектора", "группы", "центра"}
    result = []
    for w in chief_position.split():
        result.append(w)
        if w.lower() in dept_kw:
            break
    return " ".join(result)


def _strip_suffix(name: str) -> str:
    """'Мошкова Елена Владимировна (ноут)' → 'Мошкова Елена Владимировна'"""
    return re.sub(r"\s*\([^)]*\)\s*$", "", name).strip()


def _extract_phone_short(phone_str: str) -> str:
    """'8-57-4004' → '4004', '2102, 2103' → '2102', '43-29' → '4329'"""
    if not phone_str:
        return ""
    first = phone_str.split(",")[0].strip()
    # Формат 8-57-XXXX → берём последние 4 цифры после последнего дефиса
    if "-" in first:
        parts = first.split("-")
        last = re.sub(r"\D", "", parts[-1])
        if len(last) == 4:
            return last
    # Просто последние 4 цифры из строки
    digits = re.sub(r"\D", "", first)
    return digits[-4:] if len(digits) >= 4 else first


# ══════════════════════════════════════════════════════
#  Загрузка данных
# ══════════════════════════════════════════════════════


def load_phone_book(path: str) -> dict:
    wb = xlrd.open_workbook(path)
    sh = wb.sheets()[0]
    result = {}
    current_dept = ""
    dept_employees = []

    def flush_dept():
        chief = None
        for emp in dept_employees:
            pos = emp["position"].lower()
            if (
                "начальник" in pos
                and "заместитель" not in pos
                and "помощник" not in pos
            ):
                chief = emp
                break
        if chief is None and dept_employees:
            chief = dept_employees[0]
        for emp in dept_employees:
            emp["chief_name"] = chief["name"] if chief else ""
            emp["chief_position"] = chief["position"] if chief else ""
            emp["chief_initials"] = make_initials(chief["name"]) if chief else ""
            result[emp["name"]] = emp

    for i in range(sh.nrows):
        row = sh.row_values(i)
        name = str(row[0]).strip()
        pos = str(row[2]).strip() if len(row) > 2 else ""
        if not name:
            continue
        if not pos:
            flush_dept()
            dept_employees = []
            current_dept = name
        else:
            phone_raw = str(row[5]).strip() if len(row) > 5 else ""
            dept_employees.append(
                {"name": name, "position": pos, "department": current_dept, "phone": phone_raw}
            )
    flush_dept()
    return result


def load_pc_data(path: str) -> dict:
    """
    Возвращает dict: canonical_name → list[record].
    canonical_name — имя БЕЗ суффикса в скобках.
    Каждый record содержит:
      'values'     : list[str] — 11 значений колонок
      'colors'     : list[(bg,fg)] — точные #RRGGBB цвета (тема уже разрешена)
      'display_name': str — ФИО как написано в файле (с суффиксом типа "(ноут)")
      'serial', 'inventory', 'label'
    """
    theme_palette = _load_theme_from_xlsx(path)
    wb = openpyxl.load_workbook(path)
    sh = wb.active
    result = {}

    for row in sh.iter_rows(min_row=2, values_only=False):
        if len(row) < 6 or not row[5].value:
            continue
        display_name = str(row[5].value).strip()
        if not display_name:
            continue
        canonical = _strip_suffix(display_name)

        values = []
        colors = []
        for i in range(NUM_PC_COLS):
            cell = row[i] if i < len(row) else None
            if cell is None:
                values.append("")
                colors.append(("", ""))
                continue

            val = cell.value
            if isinstance(val, datetime):
                val = val.strftime("%d.%m.%Y")
            values.append(str(val).strip() if val is not None else "")

            # Фон: только если явная заливка (patternType='solid')
            has_fill = cell.fill and cell.fill.patternType == "solid"
            bg = _resolve_color(cell.fill.fgColor if has_fill else None, theme_palette)
            # #000000 и #FFFFFF — дефолтные цвета темы, не реальная заливка
            if bg in ("#000000", "#FFFFFF"):
                bg = ""

            # Текст: берём цвет если он явно задан и не чёрный по умолчанию
            fg = _resolve_color(cell.font.color if cell.font else None, theme_palette)
            if fg == "#000000":
                fg = ""  # чёрный текст = дефолт, не показываем

            colors.append((bg, fg))

        serial = values[8]
        inventory = values[9]
        cabinet = values[1]
        label = f"с/н: {serial}  |  инв: {inventory}"
        if cabinet:
            label += f"  [{cabinet}]"
        if display_name != canonical:
            label += f"  {display_name[len(canonical):]}"  # показываем суффикс "(ноут)"

        result.setdefault(canonical, []).append(
            {
                "values": values,
                "colors": colors,
                "display_name": display_name,
                "serial": serial,
                "inventory": inventory,
                "label": label,
            }
        )

    return result


def get_journal_info(path: str) -> dict:
    doc = Document(path)
    t = doc.tables[0]
    last_pp = 0
    last_reg = ""
    last_row_idx = 1
    last_executor = ""
    for i, row in enumerate(t.rows):
        vals = [c.text.strip() for c in row.cells]
        if vals[0].isdigit():
            last_pp = int(vals[0])
            last_reg = vals[2]
            last_row_idx = i
            if len(vals) > 4 and vals[4]:
                last_executor = vals[4]
    next_reg = last_reg
    if last_reg and "-" in last_reg:
        prefix, num = last_reg.rsplit("-", 1)
        try:
            next_reg = f"{prefix}-{int(num) + 1}"
        except ValueError:
            pass
    return {
        "next_pp": last_pp + 1,
        "next_reg": next_reg,
        "last_row_idx": last_row_idx,
        "last_executor": last_executor,
    }


# ══════════════════════════════════════════════════════
#  Журнал
# ══════════════════════════════════════════════════════


def _set_cell_fmt(cell, text, font_name="Times New Roman", font_size=10):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font_name)
    ex = rPr.find(qn("w:rFonts"))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, rFonts)


def add_journal_entry(path: str, entry: dict, last_row_idx: int):
    doc = Document(path)
    t = doc.tables[0]
    src_tr = t.rows[last_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
        tc.append(OxmlElement("w:p"))
    src_tr.addnext(new_tr)
    new_row = next((r for r in t.rows if r._tr is new_tr), None)
    if new_row is None:
        raise RuntimeError("Не удалось вставить строку в журнал")
    values = [
        str(entry["pp"]),
        entry["date"],
        entry["reg"],
        entry["description"],
        entry["executor"],
        entry.get("note", ""),
    ]
    for i, val in enumerate(values[: len(new_row.cells)]):
        _set_cell_fmt(new_row.cells[i], val)
    doc.save(path)


# ══════════════════════════════════════════════════════
#  Генерация Word
# ══════════════════════════════════════════════════════


def _run(
    para,
    text,
    bold=False,
    underline=False,
    italic=False,
    size=12,
    name="Times New Roman",
    superscript=False,
):
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)
    ex = rPr.find(qn("w:rFonts"))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, rFonts)
    if superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rPr.append(va)
    return run


def _para(
    container,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=0,
    first_indent_cm=None,
):
    para = container.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if first_indent_cm is not None:
        para.paragraph_format.first_line_indent = Cm(first_indent_cm)
    return para


def _set_col_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")


def _set_table_col_widths(table, widths_dxa):
    """Патчит tblGrid + ширину каждой ячейки. Без этого Word игнорирует _set_col_width."""
    tbl = table._tbl
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(grid)
    else:
        tbl.insert(0, grid)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_dxa):
                _set_col_width(cell, widths_dxa[ci])


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    ex = tblPr.find(qn("w:tblBorders"))
    if ex is not None:
        tblPr.remove(ex)
    tb = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tb.append(b)
    tblPr.append(tb)


def _cell_border(cell, sides: dict):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders")
        tcPr.append(tb)
    for side, sz in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single" if sz else "none")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tb.append(el)


def _set_tbl_cell_margins(table, left=10, right=10):
    """Устанавливает минимальные поля ячеек таблицы (как в оригинале)."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for side, val in [("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def _no_cell_borders(cell):
    _cell_border(cell, {s: 0 for s in ["top", "left", "bottom", "right"]})


def generate_zayavka(data: dict, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(0.89)
    sec.bottom_margin = Cm(2.88)

    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    fp1 = footer.paragraphs[0]
    fp1.paragraph_format.space_before = Pt(0)
    fp1.paragraph_format.space_after = Pt(0)
    _run(fp1, f'№ {data["reg_number"]}', italic=True)
    fp2 = footer.add_paragraph()
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after = Pt(0)
    _run(fp2, data["date_short"], italic=True)

    t0 = doc.add_table(rows=1, cols=3)
    _no_table_borders(t0)
    _set_col_width(t0.cell(0, 0), 3544)
    _set_col_width(t0.cell(0, 1), 1417)
    _set_col_width(t0.cell(0, 2), 3969)
    for ci in range(3):
        _no_cell_borders(t0.cell(0, ci))
    rc = t0.cell(0, 2)
    for idx, (txt, bold) in enumerate(
        [
            ("РАЗРЕШАЮ", True),
            ("Заместитель управляющего Отделением СФР по", False),
            ("Санкт-Петербургу и Ленинградской", False),
            ("области", False),
            ("_______________        Г.Г. Щемелев", False),
            ("«___» _______ 20 ___ г.", False),
        ]
    ):
        p = rc.paragraphs[0] if idx == 0 else rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    _run(p, "ЗАЯВКА", bold=True)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "НА ОБУЧЕНИЕ РАБОТЕ И УСТАНОВКУ  КРИПТОСРЕДСТВА")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, first_indent_cm=1.5)
    _run(
        p,
        "Прошу Вас разрешить подготовку к самостоятельной работе с криптографическим средством защиты информации:",
    )

    t1 = doc.add_table(rows=2, cols=1)
    _no_table_borders(t1)
    _set_col_width(t1.cell(0, 0), 9637)
    _set_col_width(t1.cell(1, 0), 9637)
    _no_cell_borders(t1.cell(0, 0))
    _cell_border(t1.cell(1, 0), {"top": 4, "bottom": 4})
    p = t1.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "СКЗИ  ViPNet CSP в составе ПО ViPNet PKI Client")
    p = t1.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(173.5)
    p.paragraph_format.first_line_indent = Pt(6.5)
    _run(p, "     (тип криптосредства)", size=8)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "работника:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_position_full"] + ",  ")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_name"])

    p = _para(doc)
    if data["with_install"]:
        _run(p, "с установкой", bold=True, underline=True)
        _run(p, " / ", italic=True)
        _run(p, "без установки", italic=True)
    else:
        _run(p, "с установкой")
        _run(p, " / ")
        _run(p, "без установки", bold=True, underline=True)
    _run(p, " криптосредства на его рабочее место.")
    p = _para(doc)
    _run(p, "                        ", size=12, superscript=True)
    _run(p, " ", size=12, superscript=True)
    _run(p, "(нужное подчеркнуть)", size=14, superscript=True)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Серийный/ инвентарный номер системного блока:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, f"с/н: {data['serial']}  инв. {data['inventory']}")

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Необходимость установки средства криптозащиты обусловлена:")
    t2 = doc.add_table(rows=1, cols=1)
    _no_table_borders(t2)
    _set_col_width(t2.cell(0, 0), 9569)
    _cell_border(t2.cell(0, 0), {"bottom": 4})
    p = t2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, data.get("purpose", "для подписания ЭП на портале ЕЦП"))
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "(наименование решаемой задачи)", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t3)
    _set_col_width(t3.cell(0, 0), 5778)
    _set_col_width(t3.cell(0, 1), 4077)
    for ci in range(2):
        _no_cell_borders(t3.cell(0, ci))
    lc = t3.cell(0, 0)
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pos_prefix = data.get("chief_pos_prefix", "")
    abbrev = data.get("chief_abbrev", "")
    if pos_prefix and abbrev:
        _run(p, pos_prefix + "             ")
        _run(p, abbrev, underline=True)
    elif abbrev:
        _run(p, abbrev, underline=True)
    else:
        _run(p, data.get("chief_position_doc", ""))
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(
        p2,
        "                                         (структурного подразделения Отделения)",
        size=9,
    )
    rc = t3.cell(0, 1)
    p = rc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "__________/", size=14)
    _run(p, data["chief_initials"], underline=True)
    p2 = rc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, "подпись", size=9)
    _run(p2, "                         ", size=12)
    _run(p2, "фамилия и инициалы", size=9)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, f"«{data['day']}» {data['month']} {data['year']} г.", underline=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = _para(doc)
    _run(p, "СОГЛАСОВАНО", bold=True)
    for line in [
        "Начальником отдела",
        "организационно-технической и ",
        "криптографической защиты информации",
    ]:
        p = _para(doc)
        _run(p, line)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, "__________________________    ")
    p = _para(doc)
    _run(p, "(подпись)", size=9)
    _run(p, "             ", size=12)
    _run(p, "(фамилия, инициалы)", size=9)
    p = _para(doc)
    _run(p, "«___» ______________ 20 ___ г.")

    doc.save(output_path)


def generate_csp_zayavka(data: dict, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(0.89)
    sec.bottom_margin = Cm(2.88)

    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    fp1 = footer.paragraphs[0]
    fp1.paragraph_format.space_before = Pt(0)
    fp1.paragraph_format.space_after = Pt(0)
    _run(fp1, f'№ {data["reg_number"]}', italic=True)
    fp2 = footer.add_paragraph()
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after = Pt(0)
    _run(fp2, data["date_short"], italic=True)

    t0 = doc.add_table(rows=1, cols=3)
    _no_table_borders(t0)
    _set_col_width(t0.cell(0, 0), 3544)
    _set_col_width(t0.cell(0, 1), 1417)
    _set_col_width(t0.cell(0, 2), 3969)
    for ci in range(3):
        _no_cell_borders(t0.cell(0, ci))
    rc = t0.cell(0, 2)
    for idx, (txt, bold) in enumerate(
        [
            ("РАЗРЕШАЮ", True),
            ("Заместитель управляющего Отделением СФР по", False),
            ("Санкт-Петербургу и Ленинградской", False),
            ("области", False),
            ("_______________        Г.Г. Щемелев", False),
            ("«___» _______ 20 ___ г.", False),
        ]
    ):
        p = rc.paragraphs[0] if idx == 0 else rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    _run(p, "ЗАЯВКА", bold=True)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "НА ОБУЧЕНИЕ РАБОТЕ И УСТАНОВКУ  КРИПТОСРЕДСТВА")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, first_indent_cm=1.5)
    _run(
        p,
        "Прошу Вас разрешить подготовку к самостоятельной работе с криптографическим средством защиты информации:",
    )

    t1 = doc.add_table(rows=2, cols=1)
    _no_table_borders(t1)
    _set_col_width(t1.cell(0, 0), 9637)
    _set_col_width(t1.cell(1, 0), 9637)
    _no_cell_borders(t1.cell(0, 0))
    _cell_border(t1.cell(1, 0), {"top": 4, "bottom": 4})
    p = t1.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "СКЗИ ViPNet CSP в составе СПО ViPNet Client")
    p = t1.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(173.5)
    p.paragraph_format.first_line_indent = Pt(6.5)
    _run(p, "     (тип криптосредства)", size=8)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "работника:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_position_full"] + ",  ")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_name"])

    p = _para(doc)
    if data["with_install"]:
        _run(p, "с установкой", bold=True, underline=True)
        _run(p, " / ", italic=True)
        _run(p, "без установки", italic=True)
    else:
        _run(p, "с установкой")
        _run(p, " / ")
        _run(p, "без установки", bold=True, underline=True)
    _run(p, " криптосредства на его рабочее место.")
    p = _para(doc)
    _run(p, "                        ", size=12, superscript=True)
    _run(p, " ", size=12, superscript=True)
    _run(p, "(нужное подчеркнуть)", size=14, superscript=True)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Серийный/ инвентарный номер системного блока:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, f"с/н: {data['serial']}  инв. {data['inventory']}")

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Необходимость установки средства криптозащиты обусловлена:")
    t2 = doc.add_table(rows=1, cols=1)
    _no_table_borders(t2)
    _set_col_width(t2.cell(0, 0), 9569)
    _cell_border(t2.cell(0, 0), {"bottom": 4})
    p = t2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(
        p,
        data.get(
            "purpose",
            "работой с ПК: ПТК КС, Элардо, АРМ БПИ, портал Казначейства, "
            "для взаимодействия по защищенному каналу по средствам VipNet "
            "«Деловая почта»",
        ),
    )
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "(наименование решаемой задачи)", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t3)
    _set_col_width(t3.cell(0, 0), 5778)
    _set_col_width(t3.cell(0, 1), 4077)
    for ci in range(2):
        _no_cell_borders(t3.cell(0, ci))
    lc = t3.cell(0, 0)
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pos_prefix = data.get("chief_pos_prefix", "")
    abbrev = data.get("chief_abbrev", "")
    if pos_prefix and abbrev:
        _run(p, pos_prefix + "             ")
        _run(p, abbrev, underline=True)
    elif abbrev:
        _run(p, abbrev, underline=True)
    else:
        _run(p, data.get("chief_position_doc", ""))
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(
        p2,
        "                                         (структурного подразделения Отделения)",
        size=9,
    )
    rc = t3.cell(0, 1)
    p = rc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "__________/", size=14)
    _run(p, data["chief_initials"], underline=True)
    p2 = rc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, "подпись", size=9)
    _run(p2, "                         ", size=12)
    _run(p2, "фамилия и инициалы", size=9)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, f"«{data['day']}» {data['month']} {data['year']} г.", underline=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = _para(doc)
    _run(p, "СОГЛАСОВАНО", bold=True)
    for line in [
        "Начальником отдела",
        "организационно-технической и ",
        "криптографической защиты информации",
    ]:
        p = _para(doc)
        _run(p, line)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, "__________________________    ")
    p = _para(doc)
    _run(p, "(подпись)", size=9)
    _run(p, "             ", size=12)
    _run(p, "(фамилия, инициалы)", size=9)
    p = _para(doc)
    _run(p, "«___» ______________ 20 ___ г.")

    doc.save(output_path)


def generate_ptk_zayavka(data: dict, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(0.89)
    sec.bottom_margin = Cm(2.88)

    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    fp1 = footer.paragraphs[0]
    fp1.paragraph_format.space_before = Pt(0)
    fp1.paragraph_format.space_after = Pt(0)
    _run(fp1, f'№ {data["reg_number"]}', italic=True)
    fp2 = footer.add_paragraph()
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after = Pt(0)
    _run(fp2, data["date_short"], italic=True)

    t0 = doc.add_table(rows=1, cols=3)
    _no_table_borders(t0)
    _set_col_width(t0.cell(0, 0), 3544)
    _set_col_width(t0.cell(0, 1), 1417)
    _set_col_width(t0.cell(0, 2), 3969)
    for ci in range(3):
        _no_cell_borders(t0.cell(0, ci))
    rc = t0.cell(0, 2)
    for idx, (txt, bold) in enumerate(
        [
            ("РАЗРЕШАЮ", True),
            ("Заместитель управляющего Отделением СФР по", False),
            ("Санкт-Петербургу и Ленинградской", False),
            ("области", False),
            ("_______________        Г.Г. Щемелев", False),
            ("«___» _______ 20 ___ г.", False),
        ]
    ):
        p = rc.paragraphs[0] if idx == 0 else rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    _run(p, "ЗАЯВКА", bold=True)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "НА ОБУЧЕНИЕ РАБОТЕ И УСТАНОВКУ  КРИПТОСРЕДСТВА")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, first_indent_cm=1.5)
    _run(
        p,
        "Прошу Вас разрешить подготовку к самостоятельной работе с криптографическим средством защиты информации:",
    )

    t1 = doc.add_table(rows=2, cols=1)
    _no_table_borders(t1)
    _set_col_width(t1.cell(0, 0), 9637)
    _set_col_width(t1.cell(1, 0), 9637)
    _no_cell_borders(t1.cell(0, 0))
    _cell_border(t1.cell(1, 0), {"top": 4, "bottom": 4})
    p = t1.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "СКЗИ ViPNet CSP")
    p = t1.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(173.5)
    p.paragraph_format.first_line_indent = Pt(6.5)
    _run(p, "     (тип криптосредства)", size=8)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "работника:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_position_full"] + ",  ")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_name"])

    p = _para(doc)
    if data["with_install"]:
        _run(p, "с установкой", bold=True, underline=True)
        _run(p, " / ", italic=True)
        _run(p, "без установки", italic=True)
    else:
        _run(p, "с установкой")
        _run(p, " / ")
        _run(p, "без установки", bold=True, underline=True)
    _run(p, " криптосредства на его рабочее место.")
    p = _para(doc)
    _run(p, "                        ", size=12, superscript=True)
    _run(p, " ", size=12, superscript=True)
    _run(p, "(нужное подчеркнуть)", size=14, superscript=True)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Серийный/ инвентарный номер системного блока:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, f"с/н: {data['serial']}  инв. {data['inventory']}")

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Необходимость установки средства криптозащиты обусловлена:")
    t2 = doc.add_table(rows=1, cols=1)
    _no_table_borders(t2)
    _set_col_width(t2.cell(0, 0), 9569)
    _cell_border(t2.cell(0, 0), {"bottom": 4})
    p = t2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, data.get("purpose", "Работой в ПТК КС"))
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "(наименование решаемой задачи)", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t3)
    _set_col_width(t3.cell(0, 0), 5778)
    _set_col_width(t3.cell(0, 1), 4077)
    for ci in range(2):
        _no_cell_borders(t3.cell(0, ci))
    lc = t3.cell(0, 0)
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pos_prefix = data.get("chief_pos_prefix", "")
    abbrev = data.get("chief_abbrev", "")
    if pos_prefix and abbrev:
        _run(p, pos_prefix + "             ")
        _run(p, abbrev, underline=True)
    elif abbrev:
        _run(p, abbrev, underline=True)
    else:
        _run(p, data.get("chief_position_doc", ""))
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(
        p2,
        "                                         (структурного подразделения Отделения)",
        size=9,
    )
    rc = t3.cell(0, 1)
    p = rc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "__________/", size=14)
    _run(p, data["chief_initials"], underline=True)
    p2 = rc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, "подпись", size=9)
    _run(p2, "                         ", size=12)
    _run(p2, "фамилия и инициалы", size=9)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, f"«{data['day']}» {data['month']} {data['year']} г.", underline=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = _para(doc)
    _run(p, "СОГЛАСОВАНО", bold=True)
    for line in [
        "Начальником отдела",
        "организационно-технической и ",
        "криптографической защиты информации",
    ]:
        p = _para(doc)
        _run(p, line)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, "__________________________    ")
    p = _para(doc)
    _run(p, "(подпись)", size=9)
    _run(p, "             ", size=12)
    _run(p, "(фамилия, инициалы)", size=9)
    p = _para(doc)
    _run(p, "«___» ______________ 20 ___ г.")

    doc.save(output_path)


def generate_kriptopro_zayavka(data: dict, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(0.89)
    sec.bottom_margin = Cm(2.88)

    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    fp1 = footer.paragraphs[0]
    fp1.paragraph_format.space_before = Pt(0)
    fp1.paragraph_format.space_after = Pt(0)
    _run(fp1, f'№ {data["reg_number"]}', italic=True)
    fp2 = footer.add_paragraph()
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after = Pt(0)
    _run(fp2, data["date_short"], italic=True)

    t0 = doc.add_table(rows=1, cols=3)
    _no_table_borders(t0)
    _set_col_width(t0.cell(0, 0), 3544)
    _set_col_width(t0.cell(0, 1), 1417)
    _set_col_width(t0.cell(0, 2), 3969)
    for ci in range(3):
        _no_cell_borders(t0.cell(0, ci))
    rc = t0.cell(0, 2)
    for idx, (txt, bold) in enumerate(
        [
            ("РАЗРЕШАЮ", True),
            ("Заместитель управляющего Отделением СФР по", False),
            ("Санкт-Петербургу и Ленинградской", False),
            ("области", False),
            ("_______________        Г.Г. Щемелев", False),
            ("«___» _______ 20 ___ г.", False),
        ]
    ):
        p = rc.paragraphs[0] if idx == 0 else rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    _run(p, "ЗАЯВКА", bold=True)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "НА ОБУЧЕНИЕ РАБОТЕ И УСТАНОВКУ  КРИПТОСРЕДСТВА")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, first_indent_cm=1.5)
    _run(
        p,
        "Прошу Вас разрешить подготовку к самостоятельной работе с криптографическим средством защиты информации:",
    )

    t1 = doc.add_table(rows=2, cols=1)
    _no_table_borders(t1)
    _set_col_width(t1.cell(0, 0), 9637)
    _set_col_width(t1.cell(1, 0), 9637)
    _no_cell_borders(t1.cell(0, 0))
    _cell_border(t1.cell(1, 0), {"top": 4, "bottom": 4})
    p = t1.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "CКЗИ КриптоПРО CSP")
    p = t1.cell(1, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(173.5)
    p.paragraph_format.first_line_indent = Pt(6.5)
    _run(p, "     (тип криптосредства)", size=8)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "работника:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_position_full"] + ",  ")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, data["employee_name"])

    p = _para(doc)
    if data["with_install"]:
        _run(p, "с установкой", bold=True, underline=True)
        _run(p, " / ", italic=True)
        _run(p, "без установки", italic=True)
    else:
        _run(p, "с установкой")
        _run(p, " / ")
        _run(p, "без установки", bold=True, underline=True)
    _run(p, " криптосредства на его рабочее место.")
    p = _para(doc)
    _run(p, "                        ", size=12, superscript=True)
    _run(p, " ", size=12, superscript=True)
    _run(p, "(нужное подчеркнуть)", size=14, superscript=True)

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Серийный/ инвентарный номер системного блока:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, f"с/н: {data['serial']}  инв. {data['inventory']}")

    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Необходимость установки средства криптозащиты обусловлена:")
    t2 = doc.add_table(rows=1, cols=1)
    _no_table_borders(t2)
    _set_col_width(t2.cell(0, 0), 9569)
    _cell_border(t2.cell(0, 0), {"bottom": 4})
    p = t2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, data.get("purpose", "работой в ЕЦП, работа на портале Казначейства"))
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "(наименование решаемой задачи)", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t3)
    _set_col_width(t3.cell(0, 0), 5778)
    _set_col_width(t3.cell(0, 1), 4077)
    for ci in range(2):
        _no_cell_borders(t3.cell(0, ci))
    lc = t3.cell(0, 0)
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pos_prefix = data.get("chief_pos_prefix", "")
    abbrev = data.get("chief_abbrev", "")
    if pos_prefix and abbrev:
        _run(p, pos_prefix + "             ")
        _run(p, abbrev, underline=True)
    elif abbrev:
        _run(p, abbrev, underline=True)
    else:
        _run(p, data.get("chief_position_doc", ""))
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(
        p2,
        "                                         (структурного подразделения Отделения)",
        size=9,
    )
    rc = t3.cell(0, 1)
    p = rc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "__________/", size=14)
    _run(p, data["chief_initials"], underline=True)
    p2 = rc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, "подпись", size=9)
    _run(p2, "                         ", size=12)
    _run(p2, "фамилия и инициалы", size=9)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, f"«{data['day']}» {data['month']} {data['year']} г.", underline=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = _para(doc)
    _run(p, "СОГЛАСОВАНО", bold=True)
    for line in [
        "Начальником отдела",
        "организационно-технической и ",
        "криптографической защиты информации",
    ]:
        p = _para(doc)
        _run(p, line)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc)
    _run(p, "__________________________    ")
    p = _para(doc)
    _run(p, "(подпись)", size=9)
    _run(p, "             ", size=12)
    _run(p, "(фамилия, инициалы)", size=9)
    p = _para(doc)
    _run(p, "«___» ______________ 20 ___ г.")

    doc.save(output_path)


# ══════════════════════════════════════════════════════
#  Акт установки PKI — генератор (с нуля, без шаблона)
# ══════════════════════════════════════════════════════


def _set_cell_margins(cell, top=0, left=108, bottom=0, right=108):
    """Индивидуальные поля ячейки."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    ex = tcPr.find(qn("w:tcMar"))
    if ex is not None:
        tcPr.remove(ex)
    mar = OxmlElement("w:tcMar")
    for side, val in [
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _akt_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False):
    """Заполнить ячейку текстом с заданным форматированием."""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    if size:
        r.font.size = Pt(size)
    r.bold = bold


def generate_ecp_zayavka(data: dict, output_path: str):
    """Генерирует заявку на предоставление доступа в ГИС ЕЦП."""
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    L = WD_ALIGN_PARAGRAPH.LEFT

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0)    # 1134 DXA — точно по оригиналу
    sec.right_margin = Cm(1.5)   # 851 DXA
    sec.top_margin = Cm(1.5)     # 851 DXA
    sec.bottom_margin = Cm(1.5)  # 851 DXA

    def _rn(para, text, sz=12, bold=False, underline=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(sz)
        r.bold = bold
        r.underline = underline
        r.italic = italic
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Times New Roman")
        ex = rPr.find(qn("w:rFonts"))
        if ex is not None:
            rPr.remove(ex)
        rPr.insert(0, rFonts)
        return r

    def _ep(container=None):
        p = (container or doc).add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    def _vAlign(cell, val="center"):
        tcPr = cell._tc.get_or_add_tcPr()
        ex = tcPr.find(qn("w:vAlign"))
        if ex is not None:
            tcPr.remove(ex)
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), val)
        tcPr.append(va)

    # ── Шапка: org слева, адресат справа ─────────────────────────
    t0 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t0)
    _set_col_width(t0.cell(0, 0), 5500)
    _set_col_width(t0.cell(0, 1), 4421)
    _no_cell_borders(t0.cell(0, 0))
    _no_cell_borders(t0.cell(0, 1))

    # Левая ячейка: организация + № + дата
    lc = t0.cell(0, 0)
    p = lc.paragraphs[0]
    p.alignment = C
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _rn(p, "Социальный фонд России", sz=9)

    for line in [
        "Отделение Фонда пенсионного и",
        "социального страхования",
        "Российской Федерации",
        "по Санкт-Петербургу и Ленинградской",
        "области",
        "(ОСФР по Санкт-Петербургу и",
        "Ленинградской области)",
    ]:
        p2 = lc.add_paragraph()
        p2.alignment = C
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        _rn(p2, line, sz=11, bold=True)

    _ep(lc)  # пустая строка

    p_num = lc.add_paragraph()
    p_num.alignment = C
    p_num.paragraph_format.space_before = Pt(0)
    p_num.paragraph_format.space_after = Pt(0)
    _rn(p_num, f'№ {data["reg_number"]}')

    p_dt = lc.add_paragraph()
    p_dt.alignment = C
    p_dt.paragraph_format.space_before = Pt(0)
    p_dt.paragraph_format.space_after = Pt(0)
    _rn(p_dt, data["date_short"])

    # Правая ячейка: адресат
    rc = t0.cell(0, 1)
    p = rc.paragraphs[0]
    p.alignment = L
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _rn(p, "Заместителю управляющего Отделением", sz=13)
    _ep(rc)
    p3 = rc.add_paragraph()
    p3.alignment = L
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    _rn(p3, "Г. Г. Щемелеву", sz=13)

    _ep()  # отступ после шапки

    # ── Заголовок ─────────────────────────────────────────────────
    p = _para(doc, align=C)
    _run(p, "ЗАЯВКА", bold=True, size=13)
    p = _para(doc, align=C)
    _run(p, "на предоставление доступа", bold=True, size=13)
    _ep()

    # ── Основной текст ────────────────────────────────────────────
    env = data.get("env", "prod")
    unlock = data.get("action_unlock", False)
    change_pass = data.get("action_change_pass", False)

    p = doc.add_paragraph()
    p.alignment = J
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.27)

    _rn(p, "Для работы в ", sz=13)
    if env == "prod":
        _rn(p, "ГИС ЕЦП (ИС КНД) Прод", sz=13, bold=True, underline=True)
        _rn(p, " / ГИС ЕЦП ТЕСТ", sz=13)
    else:
        _rn(p, "ГИС ЕЦП (ИС КНД) Прод", sz=13)
        _rn(p, " / ", sz=13)
        _rn(p, "ГИС ЕЦП ТЕСТ", sz=13, bold=True, underline=True)
    _rn(p, " прошу зарегистрировать/ изменить доступ/ добавить доступ/ отозвать доступ/ ", sz=13)
    _rn(p, "разблокировать пользователя", sz=13, bold=unlock, underline=unlock)
    _rn(p, "/ ", sz=13)
    _rn(p, "сменить пароль", sz=13, bold=change_pass, underline=change_pass)

    # (Нужное подчеркнуть, выделить) — верхняя граница вместо строки подчёркиваний
    p = _para(doc, align=C)
    _run(p, "(Нужное подчеркнуть, выделить)", bold=True)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)

    _ep()

    # ── Таблица — точные ширины из оригинала ──────────────────────
    # Col widths: 567, 1702, 1134, 1275, 1418, 1559, 1559, 1418 = 10632
    col_widths_tbl = [567, 1702, 1134, 1275, 1418, 1559, 1559, 1418]

    t1 = doc.add_table(rows=3, cols=8)
    _set_table_col_widths(t1, col_widths_tbl)

    # tblInd = -34 (как в оригинале)
    tblPr1 = t1._tbl.find(qn("w:tblPr"))
    tblInd1 = OxmlElement("w:tblInd")
    tblInd1.set(qn("w:w"), "-34")
    tblInd1.set(qn("w:type"), "dxa")
    tblPr1.append(tblInd1)

    # Строка 0: заголовки (sz=11, bold)
    hdr_lines = [
        ("№", "п/п"),
        ("Фамилия, имя, отчество", "(полностью)"),
        ("Управление, отдел", ""),
        ("Должность", ""),
        ("Шаблон роли", "ИС КНД"),
        ("Система", "ЕЦП"),
        ("Подсистема", "ЕЦП"),
        ("Шаблон роли", "ЕЦП"),
    ]
    for ci, (h1, h2) in enumerate(hdr_lines):
        cell = t1.cell(0, ci)
        _cell_border(cell, {"top": 4, "left": 4, "bottom": 4, "right": 4})
        _vAlign(cell)
        p = cell.paragraphs[0]
        p.alignment = C
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _rn(p, h1, sz=11, bold=True)
        if h2:
            p2 = cell.add_paragraph()
            p2.alignment = C
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            _rn(p2, h2, sz=11, bold=True)

    # Строка 1: номера 1-8 (sz=12)
    for ci, num in enumerate(["1", "2", "3", "4", "5", "6", "7", "8"]):
        cell = t1.cell(1, ci)
        _cell_border(cell, {"top": 4, "left": 4, "bottom": 4, "right": 4})
        _vAlign(cell)
        p = cell.paragraphs[0]
        p.alignment = C
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _rn(p, num, sz=12)

    # Строка 2: данные (sz=12)
    values = [
        ("1",                        C),
        (data.get("emp_name", ""),   L),
        (data.get("emp_dept", ""),   L),
        (data.get("emp_position",""),L),
        ("", C), ("", C), ("", C), ("", C),
    ]
    for ci, (val, align) in enumerate(values):
        cell = t1.cell(2, ci)
        _cell_border(cell, {"top": 4, "left": 4, "bottom": 4, "right": 4})
        _vAlign(cell)
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if val:
            _rn(p, val, sz=12)

    # ── Блок подписи ──────────────────────────────────────────────
    _ep()
    _ep()

    t2 = doc.add_table(rows=1, cols=2)
    _no_table_borders(t2)
    _set_col_width(t2.cell(0, 0), 7905)
    _set_col_width(t2.cell(0, 1), 2232)
    _no_cell_borders(t2.cell(0, 0))
    # Правая ячейка: только нижняя граница (линия для подписи)
    _cell_border(t2.cell(0, 1), {"bottom": 4})

    p = t2.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _rn(p, "Руководитель   структурного подразделения", sz=14)

    p = t2.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _rn(p, data.get("chief_initials", ""), sz=14)

    # Пустые строки для места подписи
    for _ in range(8):
        _ep()

    # ── Исполнитель ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _rn(p, "Исполнитель Тел", sz=14)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    phone = data.get("phone_short", "")
    _rn(p2, phone if phone else "________", sz=14)

    doc.save(output_path)

def generate_akt_pki(data: dict, output_path: str):
    """Строит Акт установки СКЗИ ViPNet PKI Client с нуля через python-docx."""
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    L = WD_ALIGN_PARAGRAPH.LEFT

    day = data["day"].zfill(2)
    month = data["month"]
    year = data["year"]
    date_str = f"«{day}» {month} {year} г."

    floor_raw = data.get("floor", "")
    # Убираем слово "этаж" если оно уже есть в строке, чтобы не было "2 этаж этаж"
    floor_clean = re.sub(r"\s*этаж\s*$", "", floor_raw, flags=re.IGNORECASE).strip()
    floor = (floor_clean + " этаж") if floor_clean else ""
    room = data.get("room", "")
    location_str = (
        f"{floor} (каб. {room}) по адресу: г. Санкт-Петербург, "
        "Сердобольская., д.64, лит.К (ул. Белоостровская, д.22)."
    )
    genitive = data["user_full_name_genitive"]
    user_name = data["user_full_name"]
    user_pos = data["user_position"]
    ootikzi_name = data["ootikzi_full_name"]
    ootikzi_pos = data["ootikzi_position"]
    # Нормализуем аббревиатуру: вставляем пробелы вокруг одиночных букв-союзов
    _abbrev_raw = data["user_dept_abbrev"]
    # 'ОПиУП' → 'ОП и УП': между заглавными группами одиночная строчная буква → ' буква '
    _abbrev = re.sub(r"([А-ЯЁ]+)([а-яё])([А-ЯЁ]+)", r"\1 \2 \3", _abbrev_raw)
    dept_label = "От " + _abbrev
    reg_number = data.get("reg_number", "")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.27)
    # Нет колонтитула

    def _ep():
        """Пустой параграф (JUSTIFY, no spacing)."""
        p = doc.add_paragraph()
        p.alignment = J
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    def _jp(text, fi_pt=None, size=12, bold=False, italic=False, underline=False):
        """Justified paragraph with optional first-line indent."""
        p = doc.add_paragraph()
        p.alignment = J
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if fi_pt is not None:
            p.paragraph_format.first_line_indent = Pt(fi_pt)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        if size:
            r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        return p

    def _cp(text, size=12, bold=False, italic=False, underline=False):
        """Center paragraph."""
        p = doc.add_paragraph()
        p.alignment = C
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        if size:
            r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        return p

    def _tb(rows, cols, widths=None):
        t = doc.add_table(rows=rows, cols=cols)
        _no_table_borders(t)
        _set_tbl_cell_margins(t)
        if widths:
            _set_table_col_widths(t, widths)
        return t

    # ── Table[0]: УТВЕРЖДАЮ ────────────────────────────────────────────────────
    # col0=6129 (пустая левая), col1=3500 (чуть уже оригинала 3847)
    t0 = _tb(1, 2)
    _set_table_col_widths(t0, [6129, 3500])
    _no_cell_borders(t0.cell(0, 0))
    _no_cell_borders(t0.cell(0, 1))

    # Левая — пустая
    lc = t0.cell(0, 0)
    lc.paragraphs[0].paragraph_format.space_before = Pt(0)
    lc.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Правая — УТВЕРЖДАЮ, индивидуальные поля как в оригинале
    rc = t0.cell(0, 1)
    _set_cell_margins(rc, top=0, left=108, bottom=0, right=108)

    # p0: УТВЕРЖДАЮ
    p0 = rc.paragraphs[0]
    p0.alignment = J
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    _run(p0, "УТВЕРЖДАЮ")

    # p1,p2: должность, fi=-11430 EMU (из оригинала), sa=0
    from docx.shared import Emu as _Emu

    for txt in [
        "Заместитель управляющего Отделением СФР по ",
        "Санкт-Петербургу и Ленинградской",
    ]:
        pp = rc.add_paragraph()
        pp.alignment = L
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.first_line_indent = _Emu(-11430)
        _run(pp, txt, size=12)

    # p3: "области"
    pp = rc.add_paragraph()
    pp.alignment = J
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after = Pt(0)
    _run(pp, "области")

    # p4: пустой
    pp = rc.add_paragraph()
    pp.alignment = J
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after = Pt(0)

    # p5: подпись
    pp = rc.add_paragraph()
    pp.alignment = J
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after = Pt(0)
    _run(pp, "______________  Г.Г. Щемелёв")

    # p6: пустой
    pp = rc.add_paragraph()
    pp.alignment = J
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after = Pt(0)

    # p7: дата с подчёркиваниями (точно как в оригинале)
    p7 = rc.add_paragraph()
    p7.alignment = J
    p7.paragraph_format.space_before = Pt(0)
    p7.paragraph_format.space_after = Pt(0)
    _run(p7, "«  ", underline=True)
    _run(p7, "    »", underline=True)
    _run(p7, " ")
    _run(p7, "                        ", underline=True)
    _run(p7, " ")
    _run(p7, "2026г", underline=True)
    _run(p7, ".")

    # Para[0,1]: два пустых параграфа
    _ep()
    _ep()

    # Para[2]: АКТ №
    _cp(f"АКТ № {reg_number}", size=12, bold=True)

    # Para[3]: заголовок
    _cp("УСТАНОВКИ И ВВОДА В ЭКСПЛУАТАЦИЮ КРИПТОСРЕДСТВ", size=12)

    # Para[4,5]: пустые
    _ep()
    _ep()

    # ── Table[1]: дата ────────────────────────────────────────────────────────
    t1 = _tb(1, 2)
    _set_col_width(t1.cell(0, 0), 4507)
    _no_cell_borders(t1.cell(0, 0))
    _set_col_width(t1.cell(0, 1), 5346)
    _no_cell_borders(t1.cell(0, 1))
    p0 = t1.cell(0, 0).paragraphs[0]
    p0.alignment = L
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    _run(p0, "г. ", italic=True)
    _run(p0, "Санкт-Петербург", italic=True)
    p1 = t1.cell(0, 1).paragraphs[0]
    p1.alignment = C
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    _run(p1, "                                              ")
    _run(p1, f"«{day}» {month} {year} г", underline=True, italic=True)
    _run(p1, ".")

    # Para[6,7]: пустые
    _ep()
    _ep()

    # ── Подписанты: без таблицы, текст с подчёркиванием ─────────────────────
    # "Мы, нижеподписавшиеся," обычный текст, затем подчёркнутый поток с ФИО
    # Исправляем 'отдел' → 'отдела' в обеих должностях
    def _fix_otdel(s):
        return re.sub(r"(?i)\bотдел\b(?!а|е|у|ом|ов)", "отдела", s)

    ootikzi_pos = _fix_otdel(ootikzi_pos)
    user_pos = _fix_otdel(user_pos)
    # Приводим должность пользователя к нижнему регистру первой буквы
    user_pos_lower = user_pos[0].lower() + user_pos[1:] if user_pos else user_pos

    p_sig = doc.add_paragraph()
    p_sig.alignment = J
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after = Pt(0)
    # "Мы, нижеподписавшиеся," — обычный, без подчёркивания
    r1 = p_sig.add_run("Мы, нижеподписавшиеся, ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    # должности + ФИО ООТиКЗИ + "и" + должность пользователя + запятая — подчёркнуто
    underlined_part = f"{ootikzi_pos} {ootikzi_name} и {user_pos_lower},"
    r2 = p_sig.add_run(underlined_part)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.italic = True
    r2.underline = True
    # Принудительный перенос строки внутри параграфа — ФИО уходит на новую строку
    br = OxmlElement("w:br")
    r2._r.append(br)
    # ФИО пользователя — НЕ подчёркнуто, на отдельной строке
    r3 = p_sig.add_run(user_name)
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(11)
    r3.italic = True

    # Para[8] — подсказка мелким шрифтом, с верхней границей (имитация линии под текстом)
    p8 = doc.add_paragraph()
    p8.alignment = C
    p8.paragraph_format.space_before = Pt(0)
    p8.paragraph_format.space_after = Pt(0)
    r8 = p8.add_run(
        "(должности, фамилии и инициалы пользователя и работника подразделения по защите информации)"
    )
    r8.font.name = "Times New Roman"
    r8.font.size = Pt(10)
    # Верхняя граница параграфа — имитирует линию под подписантами
    pPr = p8._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)

    # Para[9]
    _jp(
        "составили настоящий Акт о том, что на основании (приказа, заявки, служебной записки)",
        size=12,
    )

    # ── Table[3]: ссылка на заявку ────────────────────────────────────────────
    t3 = _tb(1, 1)
    _set_col_width(t3.cell(0, 0), 9637)
    _cell_border(t3.cell(0, 0), {"bottom": 4})
    p_ref = t3.cell(0, 0).paragraphs[0]
    p_ref.alignment = C
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after = Pt(0)
    r = p_ref.add_run(data.get("zayavka_ref", ""))
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True

    # Para[10] — подсказка мелким шрифтом
    _cp("(№,  дата документа на установку криптосредств)", size=10)

    # Para[11]: пустой JUSTIFY fi=17pt
    p11 = doc.add_paragraph()
    p11.alignment = J
    p11.paragraph_format.space_before = Pt(0)
    p11.paragraph_format.space_after = Pt(0)
    p11.paragraph_format.first_line_indent = Pt(17)

    # ── Table[4]: место установки ─────────────────────────────────────────────
    t4 = _tb(1, 2)
    _set_col_width(t4.cell(0, 0), 1590)
    _no_cell_borders(t4.cell(0, 0))
    _set_col_width(t4.cell(0, 1), 8047)
    _cell_border(t4.cell(0, 1), {"bottom": 4})
    p_loc0 = t4.cell(0, 0).paragraphs[0]
    p_loc0.alignment = J
    p_loc0.paragraph_format.space_before = Pt(0)
    p_loc0.paragraph_format.space_after = Pt(0)
    r = p_loc0.add_run("в помещении")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_loc1 = t4.cell(0, 1).paragraphs[0]
    p_loc1.alignment = J
    p_loc1.paragraph_format.space_before = Pt(0)
    p_loc1.paragraph_format.space_after = Pt(0)
    for part in [
        f"{floor} (",
        "каб",
        f". {room}) по адресу: г. Санкт-Петербург, ",
        "Сердобольская",
        "., д.64, ",
        "лит.К",
        " (ул. ",
        "Белоостровская",
        ", д.22).",
    ]:
        r = p_loc1.add_run(part)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.italic = True

    # Para[12] — подсказка мелким шрифтом
    _cp("(№ помещения, адрес)", size=10)

    # ── Table[5]: серийный номер ──────────────────────────────────────────────
    t5 = _tb(1, 2)
    _set_col_width(t5.cell(0, 0), 1080)
    _no_cell_borders(t5.cell(0, 0))
    _set_col_width(t5.cell(0, 1), 8557)
    _cell_border(t5.cell(0, 1), {"bottom": 4})
    p_s0 = t5.cell(0, 0).paragraphs[0]
    p_s0.alignment = J
    p_s0.paragraph_format.space_before = Pt(0)
    p_s0.paragraph_format.space_after = Pt(0)
    r = p_s0.add_run("на ЭВМ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_s1 = t5.cell(0, 1).paragraphs[0]
    p_s1.alignment = C
    p_s1.paragraph_format.space_before = Pt(0)
    p_s1.paragraph_format.space_after = Pt(0)
    r = p_s1.add_run(data.get("serial", ""))
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True

    # Para[13] — подсказка мелким шрифтом
    _cp("                         (№ системного блока)", size=10)

    # ── Table[6]: ФИО в родительном ───────────────────────────────────────────
    t6 = _tb(1, 2)
    _set_col_width(t6.cell(0, 0), 3000)
    _no_cell_borders(t6.cell(0, 0))
    _set_col_width(t6.cell(0, 1), 6637)
    _cell_border(t6.cell(0, 1), {"bottom": 4})
    p_g0 = t6.cell(0, 0).paragraphs[0]
    p_g0.alignment = J
    p_g0.paragraph_format.space_before = Pt(0)
    p_g0.paragraph_format.space_after = Pt(0)
    r = p_g0.add_run("находящейся в пользовании")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_g1 = t6.cell(0, 1).paragraphs[0]
    p_g1.alignment = L
    p_g1.paragraph_format.space_before = Pt(0)
    p_g1.paragraph_format.space_after = Pt(0)
    r = p_g1.add_run(genitive)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.italic = True

    # Para[14] — подсказка мелким шрифтом
    _cp(
        "                                      (фамилия и инициалы пользователя)",
        size=10,
    )

    # Para[15]: "1. Установлены:" fi=17pt, JUSTIFY
    p15 = doc.add_paragraph()
    p15.alignment = J
    p15.paragraph_format.space_before = Pt(0)
    p15.paragraph_format.space_after = Pt(0)
    p15.paragraph_format.first_line_indent = Pt(17)
    _run(p15, "1. ")
    _run(p15, "Установлены:", bold=True)

    # Para[16]: "Средство криптографической защиты информации" LEFT fi=17pt
    p16 = doc.add_paragraph()
    p16.alignment = L
    p16.paragraph_format.space_before = Pt(0)
    p16.paragraph_format.space_after = Pt(0)
    p16.paragraph_format.first_line_indent = Pt(17)
    _run(p16, "1.1.Средство криптографической защиты информации  ")

    # ── Table[7]: название СКЗИ ───────────────────────────────────────────────
    t7 = _tb(1, 1)
    _set_col_width(t7.cell(0, 0), 9637)
    _cell_border(t7.cell(0, 0), {"bottom": 4})
    p_t7 = t7.cell(0, 0).paragraphs[0]
    p_t7.alignment = C
    p_t7.paragraph_format.space_before = Pt(0)
    p_t7.paragraph_format.space_after = Pt(0)
    skzi_name = data.get(
        "skzi_name", "СКЗИ «ViPNet CSP» в составе ПО «ViPNet PKI Client»"
    )
    r = p_t7.add_run(skzi_name)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True

    # Para[17] — подсказка, размер не задаём (наследуется), как в оригинале
    p17 = doc.add_paragraph()
    p17.alignment = C
    p17.paragraph_format.space_before = Pt(0)
    p17.paragraph_format.space_after = Pt(0)
    r = p17.add_run("название ")
    r.font.name = "Times New Roman"
    r2 = p17.add_run("криптосредства")
    r2.font.name = "Times New Roman"
    r3 = p17.add_run(",")
    r3.font.name = "Times New Roman"

    # Para[18]: версии... CENTER fi=17pt sz=10
    p18 = doc.add_paragraph()
    p18.alignment = C
    p18.paragraph_format.space_before = Pt(0)
    p18.paragraph_format.space_after = Pt(0)
    p18.paragraph_format.first_line_indent = Pt(17)
    _skzi_ver = data.get("skzi_version", "1.6")
    _skzi_build = data.get("skzi_build", "1.542")
    _skzi_inv = data.get("skzi_inventory", "14/852-ОСФР")
    _run(p18, "версии\xa0", size=10)
    _run(p18, _skzi_ver, size=10, bold=True)
    _run(p18, "\xa0,\xa0сборка\xa0", size=10)
    _run(p18, _skzi_build, size=10, bold=True)
    _run(p18, "\xa0, инвентарный № ", size=10)
    _run(p18, _skzi_inv, size=10, bold=True)

    # Para[19]
    _jp("для обеспечения работы со средствами электронной подписи и шифрования.")

    # Para[20]: "1.2  Настройки..." с отступом продолжения строк
    p20 = doc.add_paragraph()
    p20.alignment = J
    p20.paragraph_format.space_before = Pt(0)
    p20.paragraph_format.space_after = Pt(0)
    p20.paragraph_format.left_indent = Pt(34)
    p20.paragraph_format.first_line_indent = Pt(-17)
    _run(
        p20,
        (
            "1.2  Настройки программного обеспечения по безопасности применения криптосредства "
            "в соответствии с правилами пользования и правами пользователя."
        ),
    )

    # ── Table[8]: Kaspersky ────────────────────────────────────────────────────
    t8 = _tb(2, 2)
    _set_col_width(t8.cell(0, 0), 5505)
    _no_cell_borders(t8.cell(0, 0))
    _set_col_width(t8.cell(0, 1), 4132)
    _cell_border(t8.cell(0, 1), {"bottom": 4})
    _no_cell_borders(t8.cell(1, 0))
    _no_cell_borders(t8.cell(1, 1))
    t8.cell(1, 0).merge(t8.cell(1, 1))

    # Сдвигаем таблицу на 17pt влево чтобы 1.3 совпало с 1.1 и 1.2
    tblPr8 = t8._tbl.find(qn("w:tblPr"))
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "340")
    tblInd.set(qn("w:type"), "dxa")
    tblPr8.append(tblInd)

    # [0,0]: "1.3  Дополнительное программное обеспечение" — перенесётся по ширине колонки
    p_k0 = t8.cell(0, 0).paragraphs[0]
    p_k0.alignment = J
    p_k0.paragraph_format.space_before = Pt(0)
    p_k0.paragraph_format.space_after = Pt(0)
    r = p_k0.add_run("1.3  Дополнительное программное обеспечение")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    # [0,1]: название AV, italic
    p_k1 = t8.cell(0, 1).paragraphs[0]
    p_k1.alignment = C
    p_k1.paragraph_format.space_before = Pt(0)
    p_k1.paragraph_format.space_after = Pt(0)
    r = p_k1.add_run("Kaspersky Endpoint Security for")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True

    # [1,0] merged: версия, italic
    _set_col_width(t8.cell(1, 0), 9637)
    p_k2 = t8.cell(1, 0).paragraphs[0]
    p_k2.alignment = L
    p_k2.paragraph_format.space_before = Pt(0)
    p_k2.paragraph_format.space_after = Pt(0)
    r = p_k2.add_run("Windows  ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True
    r2 = p_k2.add_run("11.8.0.34")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.italic = True

    # Para[21]: подсказка 8pt с верхней границей
    p21 = doc.add_paragraph()
    p21.alignment = C
    p21.paragraph_format.space_before = Pt(0)
    p21.paragraph_format.space_after = Pt(0)
    p21.paragraph_format.first_line_indent = Pt(17)
    _run(
        p21,
        " антивирусное ПО, прокси-сервер, ПО для удалённого администрирования и т.д.",
        size=8,
    )
    pPr21 = p21._p.get_or_add_pPr()
    pBdr21 = OxmlElement("w:pBdr")
    top21 = OxmlElement("w:top")
    top21.set(qn("w:val"), "single")
    top21.set(qn("w:sz"), "4")
    top21.set(qn("w:space"), "1")
    top21.set(qn("w:color"), "000000")
    pBdr21.append(top21)
    pPr21.append(pBdr21)

    # Para[22,23]: fi=17pt JUSTIFY
    _p2_suffix = data.get("p2_suffix", "")
    _jp(
        f"2. Проведена проверка целостности программного обеспечения и работоспособности криптосредства{_p2_suffix}.",
        fi_pt=17,
    )
    _jp("Установленное СПО функционирует в штатном режиме.", fi_pt=17)

    # Para[24]: пустой LEFT
    p24 = doc.add_paragraph()
    p24.alignment = L
    p24.paragraph_format.space_before = Pt(0)
    p24.paragraph_format.space_after = Pt(0)

    # ── Table[9]: обучение ────────────────────────────────────────────────────
    t9 = _tb(2, 2)
    _set_col_width(t9.cell(0, 0), 4819)
    _no_cell_borders(t9.cell(0, 0))
    _set_col_width(t9.cell(0, 1), 4818)
    _cell_border(t9.cell(0, 1), {"bottom": 4})
    _no_cell_borders(t9.cell(1, 0))
    _no_cell_borders(t9.cell(1, 1))
    p_tr0 = t9.cell(0, 0).paragraphs[0]
    p_tr0.alignment = J
    p_tr0.paragraph_format.space_before = Pt(0)
    p_tr0.paragraph_format.space_after = Pt(0)
    r = p_tr0.add_run("3.Проведено обучение с принятием зачета")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_tr1 = t9.cell(0, 1).paragraphs[0]
    p_tr1.alignment = C
    p_tr1.paragraph_format.space_before = Pt(0)
    p_tr1.paragraph_format.space_after = Pt(0)
    _p3_genitive_prefix = data.get("p3_genitive_prefix", "")
    if _p3_genitive_prefix:
        rp = p_tr1.add_run(_p3_genitive_prefix)
        rp.font.name = "Times New Roman"
        rp.font.size = Pt(11)
        rp.italic = True
    r = p_tr1.add_run(genitive)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.italic = True
    p_tr2 = t9.cell(1, 1).paragraphs[0]
    p_tr2.alignment = C
    p_tr2.paragraph_format.space_before = Pt(0)
    p_tr2.paragraph_format.space_after = Pt(0)
    r = p_tr2.add_run("(фамилия и инициалы пользователя)")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # Para[25]: JUSTIFY (длинный текст без fi)
    _jp(
        (
            "по правилам пользования, соблюдению требований безопасности при работе с "
            "криптосредствами и выполнению правил обращения с ключевыми носителями, "
            "а также об ответственности за их нарушение. Вопросов нет."
        )
    )

    # Para[26]: "4. Помещение ... адрес ..." JUSTIFY fi=17pt, частично italic
    p26 = doc.add_paragraph()
    p26.alignment = J
    p26.paragraph_format.space_before = Pt(0)
    p26.paragraph_format.space_after = Pt(0)
    p26.paragraph_format.first_line_indent = Pt(17)
    _run(p26, "4. Помещение ")
    _run(
        p26,
        (
            f"{floor} (каб. {room})    по адресу: г. Санкт-Петербург, "
            "Сердобольская., д.64, лит.К (ул. Белоостровская, д.22) "
        ),
        italic=True,
    )
    _run(
        p26,
        (
            "и его оборудование, размещение ЭВМ с установленным программным "
            "(программно-аппаратным) криптосредством, хранилища ключевых документов, "
            "охрана помещения и подготовленность пользователя к самостоятельной "
            "эксплуатации СКЗ  соответствуют требованиям Инструкции по организации "
            "криптографической защиты информации в Пенсионном фонде Российской Федерации."
        ),
    )

    # Para[27,28,29]: fi=17pt JUSTIFY
    _jp("5. Акт подготовлен в 1 экземпляре:", fi_pt=17)
    _jp(
        "-    экземпляр – ООТ и КЗИ ОСФР по Санкт-Петербургу и Ленинградской области;",
        fi_pt=34,
    )
    # Para[29]: подсказка мелким шрифтом — по центру под "экземпляр"
    p29 = doc.add_paragraph()
    p29.alignment = C
    p29.paragraph_format.space_before = Pt(0)
    p29.paragraph_format.space_after = Pt(0)
    r = p29.add_run("(наименование подразделения пользователя)")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # ── Table[10]: ПАК ────────────────────────────────────────────────────────
    t10 = _tb(2, 2)
    tblPr10 = t10._tbl.find(qn("w:tblPr"))
    tblInd10 = OxmlElement("w:tblInd")
    tblInd10.set(qn("w:w"), "340")
    tblInd10.set(qn("w:type"), "dxa")
    tblPr10.append(tblInd10)
    _set_col_width(t10.cell(0, 0), 4365)
    _no_cell_borders(t10.cell(0, 0))
    _set_col_width(t10.cell(0, 1), 5272)
    _cell_border(t10.cell(0, 1), {"bottom": 4})
    _no_cell_borders(t10.cell(1, 0))
    _no_cell_borders(t10.cell(1, 1))
    p_pk0 = t10.cell(0, 0).paragraphs[0]
    p_pk0.alignment = J
    p_pk0.paragraph_format.space_before = Pt(0)
    p_pk0.paragraph_format.space_after = Pt(0)
    r = p_pk0.add_run("6. Программно-аппаратный комплекс")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p_pk1 = t10.cell(0, 1).paragraphs[0]
    _pak_align = C if data.get("pak_center") else J
    p_pk1.alignment = _pak_align
    p_pk1.paragraph_format.space_before = Pt(0)
    p_pk1.paragraph_format.space_after = Pt(0)
    r = p_pk1.add_run(skzi_name)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True
    p_pk2 = t10.cell(1, 1).paragraphs[0]
    p_pk2.alignment = C
    p_pk2.paragraph_format.space_before = Pt(0)
    p_pk2.paragraph_format.space_after = Pt(0)
    r = p_pk2.add_run("(тип, наименование криптосредства)")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # Para[30]: пустой
    _ep()

    # ── Table[11]: пломба-наклейка ────────────────────────────────────────────
    t11 = _tb(2, 4)
    for ri in range(2):
        for ci in range(4):
            _no_cell_borders(t11.cell(ri, ci))
    for ci, w in enumerate([2385, 2433, 1422, 3397]):
        _set_col_width(t11.cell(0, ci), w)
        _set_col_width(t11.cell(1, ci), w)
    _cell_border(t11.cell(0, 1), {"bottom": 4})
    _cell_border(t11.cell(0, 3), {"bottom": 4})
    # border [1,1] убран (номер печати без границы)
    # border [1,3] убран (фамилия и инициалы пользователя без границы)

    p_s0 = t11.cell(0, 0).paragraphs[0]
    p_s0.alignment = J
    p_s0.paragraph_format.space_before = Pt(0)
    p_s0.paragraph_format.space_after = Pt(0)
    r = p_s0.add_run("опечатанный пломбой-наклейкой")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    p_sn = t11.cell(0, 1).paragraphs[0]
    p_sn.alignment = C
    p_sn.paragraph_format.space_before = Pt(14)
    p_sn.paragraph_format.space_after = Pt(0)
    r = p_sn.add_run(data.get("sticker_number", ""))
    r.font.name = "Times New Roman"
    r.italic = True

    p_pr = t11.cell(0, 2).paragraphs[0]
    p_pr.alignment = C
    p_pr.paragraph_format.space_before = Pt(14)
    p_pr.paragraph_format.space_after = Pt(0)
    r = p_pr.add_run("принят")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    p_ui = t11.cell(0, 3).paragraphs[0]
    p_ui.alignment = C
    p_ui.paragraph_format.space_before = Pt(14)
    p_ui.paragraph_format.space_after = Pt(0)
    r = p_ui.add_run(data["user_initials"])
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.italic = True

    p_np = t11.cell(1, 1).paragraphs[0]
    p_np.alignment = C
    p_np.paragraph_format.space_before = Pt(0)
    p_np.paragraph_format.space_after = Pt(0)
    r = p_np.add_run("(номер печати)")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    p_nf = t11.cell(1, 3).paragraphs[0]
    p_nf.alignment = C
    p_nf.paragraph_format.space_before = Pt(0)
    p_nf.paragraph_format.space_after = Pt(0)
    r = p_nf.add_run("(фамилия и инициалы пользователя)")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # Para[31]: пустой
    _ep()

    # Para[32]: "Замечаний нет." LEFT sz=12
    p32 = doc.add_paragraph()
    p32.alignment = L
    p32.paragraph_format.space_before = Pt(0)
    p32.paragraph_format.space_after = Pt(0)
    _run(p32, "Замечаний нет.")

    # Para[33]: пустой
    _ep()

    # ── Table[12]: подписи ────────────────────────────────────────────────────
    # 6 строк × 5 столбцов, ширины: [2220, 2040, 1185, 2096, 2096]
    t12 = _tb(6, 5)
    col_widths = [2220, 2040, 1185, 2096, 2096]
    for ri in range(6):
        for ci in range(5):
            _no_cell_borders(t12.cell(ri, ci))
            _set_col_width(t12.cell(ri, ci), col_widths[ci])

    def _sig_cell(r, c, text, align=L, sz=None, italic=False):
        p = t12.cell(r, c).paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        rn = p.add_run(text)
        rn.font.name = "Times New Roman"
        if sz:
            rn.font.size = Pt(sz)
        rn.italic = italic

    # Строка 0: "От " обычный + аббревиатура italic — с отступом сверху
    t12.cell(0, 0).merge(t12.cell(0, 1))
    t12.cell(0, 3).merge(t12.cell(0, 4))
    for cell, abbrev, aln in [
        (t12.cell(0, 0), "ООТ и КЗИ", L),
        (t12.cell(0, 3), dept_label.replace("От ", ""), C),
    ]:
        p = cell.paragraphs[0]
        p.alignment = aln
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run("От ")
        r1.font.name = "Times New Roman"
        r2 = p.add_run(abbrev)
        r2.font.name = "Times New Roman"
        r2.italic = True

    # Строка 1: captions size=10
    t12.cell(1, 0).merge(t12.cell(1, 1))
    t12.cell(1, 3).merge(t12.cell(1, 4))
    _sig_cell(1, 0, "(наименование подразделения)", L, sz=10)
    _sig_cell(1, 3, "(наименование подразделения)", C, sz=10)

    # Строка 2: пустая — без границ, просто отступ для области подписи
    t12.cell(2, 1).paragraphs[0].paragraph_format.space_before = Pt(20)
    t12.cell(2, 4).paragraphs[0].paragraph_format.space_before = Pt(20)

    # Строка 3: инициалы italic — граница СНИЗУ (подпись под ФИО)
    _sig_cell(3, 1, data["ootikzi_initials"], C, italic=True)
    _sig_cell(3, 4, data["user_initials"], C, italic=True)
    _cell_border(t12.cell(3, 0), {"bottom": 4})  # левая ячейка от Шеляпин
    _cell_border(t12.cell(3, 1), {"bottom": 4})
    _cell_border(t12.cell(3, 3), {"bottom": 4})  # левая ячейка от Немчинова
    _cell_border(t12.cell(3, 4), {"bottom": 4})
    t12.cell(3, 1).paragraphs[0].paragraph_format.space_before = Pt(4)
    t12.cell(3, 4).paragraphs[0].paragraph_format.space_before = Pt(4)

    # Строка 4: captions size=10
    t12.cell(4, 0).merge(t12.cell(4, 1))
    t12.cell(4, 3).merge(t12.cell(4, 4))
    _sig_cell(
        4,
        0,
        "(подпись и фамилия работника установившего криптосредство и выдавшего документы)",
        C,
        sz=10,
    )
    _sig_cell(4, 3, "(подпись  и фамилия пользователя)", C, sz=10)

    # Строка 5: дата italic
    t12.cell(5, 0).merge(t12.cell(5, 1))
    t12.cell(5, 3).merge(t12.cell(5, 4))
    _sig_cell(5, 0, date_str, L, sz=12, italic=True)
    _sig_cell(5, 3, date_str, C, sz=12, italic=True)

    # Para[34]: пустой
    _ep()

    doc.save(output_path)


# ══════════════════════════════════════════════════════
#  Lock-файл для защиты журнала от одновременной записи
# ══════════════════════════════════════════════════════

LOCK_TTL = 60  # секунд, после которых lock считается устаревшим


def _journal_lock_path(journal_path: str) -> str:
    return journal_path + ".lock"


def _read_lock_info(lock_path: str) -> str:
    """Возвращает строку 'ФИО (ИМЯ_ПК)' из lock-файла."""
    try:
        with open(lock_path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        hostname = lines[0].strip() if lines else "неизвестный"
        fio = lines[2].strip() if len(lines) > 2 else ""
        return f"{fio} ({hostname})" if fio else hostname
    except Exception:
        return "неизвестный"


def _lock_age_seconds(lock_path: str) -> float:
    try:
        mtime = os.path.getmtime(lock_path)
        return (datetime.now() - datetime.fromtimestamp(mtime)).total_seconds()
    except Exception:
        return LOCK_TTL + 1  # считаем устаревшим если не смогли прочитать


def acquire_journal_lock(journal_path: str, executor_name: str = "") -> tuple:
    """Пытается захватить блокировку журнала.
    Возвращает (True, '') при успехе,
               (False, 'ФИО (ИМЯ_ПК)') если заблокировано другим пользователем,
               (False, None) если не удалось создать lock-файл."""
    lock_path = _journal_lock_path(journal_path)
    if os.path.exists(lock_path):
        age = _lock_age_seconds(lock_path)
        if age < LOCK_TTL:
            return False, _read_lock_info(lock_path)
        # Lock устарел — перезахватываем
    try:
        with open(lock_path, "w", encoding="utf-8") as f:
            f.write(
                f"{socket.gethostname()}\n{datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n{executor_name}"
            )
        return True, ""
    except Exception:
        return False, None  # Не удалось создать lock-файл (нет прав?)


def release_journal_lock(journal_path: str):
    """Удаляет lock-файл журнала."""
    lock_path = _journal_lock_path(journal_path)
    try:
        if os.path.exists(lock_path):
            os.remove(lock_path)
    except Exception:
        pass


def _write_journal_with_retry(journal_path: str, entry: dict, last_row_idx: int) -> str:
    """Записывает строку в журнал. При PermissionError (файл открыт в Word на Windows)
    показывает диалог с кнопками Повторить / Пропустить.
    Возвращает строку статуса для показа пользователю."""
    while True:
        try:
            add_journal_entry(journal_path, entry, last_row_idx)
            return f"✓ Запись №{entry['pp']} добавлена в журнал."
        except PermissionError:
            retry = messagebox.askretrycancel(
                "Журнал открыт в Word",
                "Не удалось записать в журнал — файл занят программой Word.\n\n"
                "Закройте журнал в Word и нажмите «Повторить».\n"
                "Нажмите «Отмена», чтобы пропустить запись в журнал.",
            )
            if not retry:
                return "⚠ Журнал не обновлён — файл был открыт в Word."
            # иначе — повторяем попытку
        except Exception as e:
            return f"⚠ Журнал не обновлён: {e}"


# ══════════════════════════════════════════════════════
#  Проверка дубликатов в журнале
# ══════════════════════════════════════════════════════


def check_journal_duplicate(journal_path: str, emp_initials: str, serial: str):
    """Ищет запись где в описании (col 3) одновременно есть инициалы и серийный номер.
    Возвращает строку вида 'Запись №5 от 15.01.2026' или None."""
    try:
        doc = Document(journal_path)
        t = doc.tables[0]
        for row in t.rows:
            vals = [c.text.strip() for c in row.cells]
            if len(vals) > 3:
                desc = vals[3]
                if emp_initials and emp_initials in desc:
                    if serial and serial in desc:
                        pp = vals[0] if vals[0] else "?"
                        date = vals[1] if len(vals) > 1 else ""
                        return f"Запись №{pp} от {date}"
    except Exception:
        pass
    return None


def search_journal_entries(journal_path: str, query: str) -> list:
    """Ищет в журнале строки, где в описании (col 3) есть query.
    Возвращает список dict: pp, date, reg, desc, executor, ref (до 20 шт.)."""
    results = []
    try:
        doc = Document(journal_path)
        t = doc.tables[0]
        q = query.strip().lower()
        for row in t.rows:
            vals = [c.text.strip() for c in row.cells]
            if len(vals) > 3 and (
                not q
                or q in vals[3].lower()
                or (len(vals) > 4 and q in vals[4].lower())
            ):
                reg = vals[2] if len(vals) > 2 else ""
                date = vals[1] if len(vals) > 1 else ""
                results.append(
                    {
                        "pp": vals[0],
                        "date": date,
                        "reg": reg,
                        "desc": vals[3],
                        "executor": vals[4] if len(vals) > 4 else "",
                        "ref": f"заявки {reg} от {date} г" if reg else vals[3][:60],
                    }
                )
    except Exception:
        pass
    return results


def _to_genitive(full_name: str) -> str:
    """Преобразует ФИО из именительного в родительный падеж.
    Пол определяется по окончанию отчества."""
    parts = full_name.strip().split()
    if len(parts) != 3:
        return full_name
    surname, name, patr = parts
    pl = patr.lower()
    if pl.endswith(("овна", "евна", "ична")):
        gender = "f"
    elif pl.endswith(("ович", "евич")):
        gender = "m"
    else:
        return full_name

    vowels = set("аеёиоуыюяАЕЁИОУЫЮЯ")

    def _f_surname(s):
        sl = s.lower()
        if sl.endswith("ова"):
            return s[:-1] + "ой"
        if sl.endswith("ева"):
            return s[:-1] + "ой"
        if sl.endswith("ина"):
            return s[:-1] + "ой"
        if sl.endswith("ая"):
            return s[:-2] + "ой"
        if sl.endswith("яя"):
            return s[:-2] + "ей"
        if sl.endswith("а"):
            return s[:-1] + "ы"
        return s

    def _m_surname(s):
        sl = s.lower()
        if sl.endswith("ов"):
            return s + "а"
        if sl.endswith("ев"):
            return s + "а"
        if sl.endswith("ин"):
            return s + "а"
        if sl.endswith("ий"):
            return s[:-2] + "ого"
        if sl.endswith("ый"):
            return s[:-2] + "ого"
        if sl.endswith("ь"):
            return s[:-1] + "я"
        if s[-1] not in vowels:
            return s + "а"
        return s

    def _f_name(s):
        sl = s.lower()
        if sl.endswith("ья"):
            return s[:-2] + "ьи"
        if sl.endswith("ия"):
            return s[:-1] + "и"
        if sl.endswith("я"):
            return s[:-1] + "и"
        if sl.endswith("а"):
            return s[:-1] + "ы"
        return s

    def _m_name(s):
        sl = s.lower()
        if sl.endswith("ий"):
            return s[:-2] + "ия"
        if sl.endswith("й"):
            return s[:-1] + "я"
        if sl.endswith("ь"):
            return s[:-1] + "я"
        if s[-1] not in vowels:
            return s + "а"
        return s

    def _f_patr(s):
        sl = s.lower()
        if sl.endswith("овна"):
            return s[:-1] + "ы"
        if sl.endswith("евна"):
            return s[:-1] + "ы"
        if sl.endswith("ична"):
            return s[:-1] + "ы"
        return s

    def _m_patr(s):
        sl = s.lower()
        if sl.endswith("ович"):
            return s + "а"
        if sl.endswith("евич"):
            return s + "а"
        if sl.endswith("ич"):
            return s + "а"
        return s

    if gender == "f":
        return f"{_f_surname(surname)} {_f_name(name)} {_f_patr(patr)}"
    else:
        return f"{_m_surname(surname)} {_m_name(name)} {_m_patr(patr)}"


# ══════════════════════════════════════════════════════
#  GUI — Excel Preview Table
# ══════════════════════════════════════════════════════


class ExcelPreviewTable(tk.Frame):
    FONT = ("TkDefaultFont", 11)
    FONT_BOLD = ("TkDefaultFont", 11, "bold")
    HDR_BG = "#2b5797"
    HDR_FG = "#ffffff"
    EMPTY_BG = "#f5f5f5"
    ROW_H = 30
    COL_W = [70, 90, 145, 170, 110, 185, 108, 170, 160, 160, 90, 36]

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self._cells = []  # list of tk.Label (data cells)
        self._build()

    def _build(self):
        self._hbar = ttk.Scrollbar(self, orient="horizontal")
        self._hbar.pack(side="bottom", fill="x")
        self._canvas = tk.Canvas(
            self,
            height=self.ROW_H * 2 + 4,
            bd=0,
            highlightthickness=0,
            xscrollcommand=self._hbar.set,
            xscrollincrement=30,
        )
        self._canvas.pack(side="top", fill="both", expand=True)
        self._hbar.config(command=self._canvas.xview)

        self._inner = tk.Frame(self._canvas, bg="#cccccc")
        self._win = self._canvas.create_window((0, 0), window=self._inner, anchor="nw")
        self._inner.bind(
            "<Configure>",
            lambda _: self._canvas.configure(scrollregion=self._canvas.bbox("all")),
        )
        self._canvas.bind(
            "<Configure>",
            lambda _: self._canvas.configure(scrollregion=self._canvas.bbox("all")),
        )

        def _on_hscroll(e):
            d = (
                int(-1 * (e.delta / 120))
                if abs(e.delta) >= 120
                else (-1 if e.delta > 0 else 1)
            )
            self._canvas.xview_scroll(d, "units")
            return "break"  # не передавать событие дальше в bind_all

        # Заголовки
        for ci, h in enumerate(PC_HEADERS):
            w = self.COL_W[ci] if ci < len(self.COL_W) else 90
            fr = tk.Frame(self._inner, width=w, height=self.ROW_H, bg="#cccccc")
            fr.grid(row=0, column=ci, padx=(0, 1), pady=(0, 1))
            fr.grid_propagate(False)
            lbl = tk.Label(
                fr,
                text=h,
                font=self.FONT_BOLD,
                bg=self.HDR_BG,
                fg=self.HDR_FG,
                anchor="center",
                padx=4,
                wraplength=w - 8,
                justify="center",
            )
            lbl.place(relwidth=1, relheight=1)

        self._show_empty()

    def _show_empty(self):
        for lbl in self._cells:
            try:
                lbl.master.destroy()
            except:
                pass
        self._cells = []
        for ci in range(NUM_PC_COLS):
            w = self.COL_W[ci] if ci < len(self.COL_W) else 90
            fr = tk.Frame(self._inner, width=w, height=self.ROW_H, bg="#cccccc")
            fr.grid(row=1, column=ci, padx=(0, 1), pady=(0, 1))
            fr.grid_propagate(False)
            lbl = tk.Label(
                fr,
                text="—",
                font=self.FONT,
                bg=self.EMPTY_BG,
                fg="#999999",
                anchor="center",
            )
            lbl.place(relwidth=1, relheight=1)
            self._cells.append(lbl)
        self._canvas.configure(height=self.ROW_H * 2 + 6)

    def show(self, pc_records: list):
        # Удаляем старые строки данных
        for lbl in self._cells:
            try:
                lbl.master.destroy()
            except:
                pass
        self._cells = []

        n = len(pc_records)
        for ri, rec in enumerate(pc_records):
            row_idx = ri + 1
            alt_bg = "#f9f9f9" if ri % 2 == 0 else "#ffffff"
            for ci in range(NUM_PC_COLS):
                val = rec["values"][ci] if ci < len(rec["values"]) else ""
                raw_bg, raw_fg = (
                    rec["colors"][ci] if ci < len(rec["colors"]) else ("", "")
                )
                bg = raw_bg if raw_bg else alt_bg
                fg = _readable_fg(bg, raw_fg)

                w = self.COL_W[ci] if ci < len(self.COL_W) else 90
                fr = tk.Frame(self._inner, width=w, height=self.ROW_H, bg="#cccccc")
                fr.grid(row=row_idx, column=ci, padx=(0, 1), pady=(0, 1))
                fr.grid_propagate(False)
                lbl = tk.Label(
                    fr,
                    text=val,
                    font=self.FONT,
                    bg=bg,
                    fg=fg,
                    anchor="w",
                    padx=4,
                    wraplength=w - 8,
                    justify="left",
                )
                lbl.place(relwidth=1, relheight=1)
                self._cells.append(lbl)

        visible_rows = min(n, 5)
        self._canvas.configure(height=self.ROW_H * (visible_rows + 1) + 6)

    def clear(self):
        self._show_empty()


# ══════════════════════════════════════════════════════
#  GUI — BaseZayavkaApp
# ══════════════════════════════════════════════════════


class BaseZayavkaApp(tk.Toplevel):
    PURPOSE_OPTIONS = []  # переопределяется в дочерних классах

    def __init__(self, launcher):
        super().__init__(launcher)
        self._launcher = launcher
        self.geometry("800x1000")
        self.resizable(True, True)
        self.cfg = load_config()
        self.phone_data = {}
        self.pc_data = {}
        self.journal_info = {}
        self.all_names = []
        self._pcs = []
        self._chief_pos_prefix = ""
        self._chief_abbrev = ""
        self.protocol("WM_DELETE_WINDOW", self._back_to_launcher)
        self._build_ui()
        self._load_files()
        self._bind_clipboard()

    def _back_to_launcher(self):
        self._launcher.deiconify()
        self.destroy()

    def _bind_clipboard(self):
        pass  # clipboard handled at root level in LauncherWindow

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        raise NotImplementedError

    def _generate_doc(self, data: dict, output_path: str):
        raise NotImplementedError

    def _journal_description(self, emp_initials: str, serial: str) -> str:
        raise NotImplementedError

    def _build_ui(self):
        # Главный canvas + scrollbar для вертикального скролла всего окна
        main_canvas = tk.Canvas(
            self, borderwidth=0, highlightthickness=0, yscrollincrement=20
        )
        vbar = ttk.Scrollbar(self, orient="vertical", command=main_canvas.yview)
        main_canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        main_canvas.pack(side="left", fill="both", expand=True)

        main = tk.Frame(main_canvas)
        self._main_win = main_canvas.create_window((0, 0), window=main, anchor="nw")
        main.bind(
            "<Configure>",
            lambda _: main_canvas.configure(scrollregion=main_canvas.bbox("all")),
        )
        main_canvas.bind(
            "<Configure>",
            lambda e: main_canvas.itemconfig(self._main_win, width=e.width),
        )

        # Скролл колёсиком — всегда главная страница
        def _scroll_delta(event):
            return int(-1 * (event.delta / 120)) if abs(event.delta) >= 120 else (-1 if event.delta > 0 else 1)

        def _on_mousewheel(event):
            if not event.delta:
                return
            if not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(_scroll_delta(event), 'units')

        def _on_wheel_linux(event, direction):
            if not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(direction, 'units')

        self.bind_all('<MouseWheel>', _on_mousewheel)
        self.bind_all('<Button-4>', lambda e: _on_wheel_linux(e, -1))
        self.bind_all('<Button-5>', lambda e: _on_wheel_linux(e,  1))

        PAD = dict(padx=8, pady=4)

        # ── Назад ─────────────────────────────────────
        ttk.Button(main, text="← Назад", command=self._back_to_launcher).pack(
            anchor="w", padx=8, pady=(6, 0)
        )

        # ── Файлы ────────────────────────────────────
        ff = ttk.LabelFrame(main, text="  Файлы данных  ", padding=8)
        ff.pack(fill="x", **PAD)
        file_specs = [
            (
                "phone_book",
                "Телефонный справочник (.xls):",
                [("XLS", "*.xls"), ("Все", "*.*")],
            ),
            (
                "pc_file",
                "Актуализация ПК (.xlsx):",
                [("XLSX", "*.xlsx"), ("Все", "*.*")],
            ),
            (
                "journal",
                "Журнал регистрации (.docx):",
                [("DOCX", "*.docx"), ("Все", "*.*")],
            ),
        ]
        self._file_vars = {}
        for ri, (key, label, ftypes) in enumerate(file_specs):
            ttk.Label(ff, text=label, width=30, anchor="w").grid(
                row=ri, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=self.cfg.get(key, ""))
            self._file_vars[key] = var
            ttk.Entry(ff, textvariable=var, width=42).grid(row=ri, column=1, padx=4)
            ttk.Button(
                ff,
                text="…",
                width=3,
                command=lambda k=key, ft=ftypes: self._pick_file(k, ft),
            ).grid(row=ri, column=2)
        ttk.Button(
            ff, text="↺  Загрузить / обновить файлы", command=self._load_files
        ).grid(row=3, column=0, columnspan=3, pady=(6, 0))

        # ── Поиск + статус ────────────────────────────
        fs = ttk.LabelFrame(main, text="  Поиск сотрудника  ", padding=8)
        fs.pack(fill="x", **PAD)
        fs.columnconfigure(1, weight=1)

        # Строка поиска
        ttk.Label(fs, text="ФИО:").grid(row=0, column=0, sticky="w")
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self._on_search_change)
        ttk.Entry(fs, textvariable=self.search_var, width=40).grid(
            row=0, column=1, padx=4, sticky="ew"
        )
        ttk.Button(fs, text="Найти", command=self._do_search).grid(row=0, column=2)

        # Статус сразу под строкой поиска (в области поиска, справа от списка)
        self.status_var = tk.StringVar(value="")
        self._status_lbl = tk.Label(
            fs,
            textvariable=self.status_var,
            foreground="#555",
            wraplength=550,
            anchor="w",
            justify="left",
            font=("TkDefaultFont", 11, "bold"),
        )
        self._status_lbl.grid(row=0, column=3, padx=(12, 4), sticky="ew")
        fs.columnconfigure(3, weight=1)

        # Список результатов
        self.lb = tk.Listbox(fs, height=10, width=50, font=("TkDefaultFont", 10))
        self.lb.grid(row=1, column=0, columnspan=2, pady=(4, 0), sticky="ew")
        self.lb.bind("<<ListboxSelect>>", self._on_select)
        sb = ttk.Scrollbar(fs, orient="vertical", command=self.lb.yview)
        sb.grid(row=1, column=2, sticky="ns", pady=(4, 0))
        self.lb.configure(yscrollcommand=sb.set)
        self.lb.bind('<MouseWheel>', lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(int(-1*(e.delta/120)) if abs(e.delta)>=120 else (-1 if e.delta>0 else 1), 'units'), 'break')[1] if e.delta else 'break')
        self.lb.bind('<Button-4>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(-1, 'units'), 'break')[1])
        self.lb.bind('<Button-5>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll( 1, 'units'), 'break')[1])

        # ── Excel-превью ─────────────────────────────
        fp = ttk.LabelFrame(main, text="  Данные из таблицы актуализации  ", padding=6)
        fp.pack(fill="x", **PAD)
        self._preview = ExcelPreviewTable(fp)
        self._preview.pack(fill="x", expand=True)
        self._cert_var = tk.StringVar()
        self._cert_lbl = tk.Label(
            fp,
            textvariable=self._cert_var,
            anchor="w",
            font=("TkDefaultFont", 10, "bold"),
        )
        self._cert_lbl.pack(fill="x", padx=4, pady=(4, 0))

        # ── Данные сотрудника ─────────────────────────
        fe = ttk.LabelFrame(
            main, text="  Данные сотрудника (для документа)  ", padding=8
        )
        fe.pack(fill="x", **PAD)
        fe.columnconfigure(1, weight=1)
        self._emp = {}
        for i, (key, label) in enumerate(
            [
                ("emp_name", "ФИО сотрудника:"),
                ("emp_position_doc", "Должность и отдел (полная строка):"),
                ("chief_name", "Начальник (ФИО):"),
                ("chief_initials", "Инициалы начальника (для подписи):"),
                ("chief_position_doc", "Аббревиатура отдела (для подписи):"),
            ]
        ):
            ttk.Label(fe, text=label, anchor="w", width=38).grid(
                row=i, column=0, sticky="w", pady=2
            )
            var = tk.StringVar()
            self._emp[key] = var
            ttk.Entry(fe, textvariable=var, width=46).grid(
                row=i, column=1, padx=4, sticky="ew"
            )

        # ── Компьютер ─────────────────────────────────
        fpc = ttk.LabelFrame(main, text="  Компьютер  ", padding=8)
        fpc.pack(fill="x", **PAD)
        fpc.columnconfigure(1, weight=1)
        ttk.Label(fpc, text="Выбор ПК:", anchor="w", width=22).grid(
            row=0, column=0, sticky="w", pady=2
        )
        self.pc_combo = ttk.Combobox(fpc, width=50, state="readonly")
        self.pc_combo.grid(row=0, column=1, padx=4, sticky="ew")
        self.pc_combo.bind("<<ComboboxSelected>>", self._on_pc_select)
        self._serial_var = tk.StringVar()
        self._inv_var = tk.StringVar()
        for ri, (txt, var) in enumerate(
            [
                ("Серийный номер:", self._serial_var),
                ("Инвентарный номер:", self._inv_var),
            ],
            2,
        ):
            ttk.Label(fpc, text=txt, anchor="w", width=22).grid(
                row=ri, column=0, sticky="w", pady=2
            )
            ttk.Entry(fpc, textvariable=var, width=40).grid(
                row=ri, column=1, padx=4, sticky="ew"
            )

        # ── Цель установки ────────────────────────────
        fpur = ttk.LabelFrame(
            main, text="  Необходимость установки криптозащиты обусловлена  ", padding=8
        )
        fpur.pack(fill="x", **PAD)
        fpur.columnconfigure(1, weight=1)
        ttk.Label(fpur, text="Выбор:", anchor="w", width=10).grid(
            row=0, column=0, sticky="w", pady=2
        )
        self._purpose_combo_var = tk.StringVar()
        self._purpose_combo = ttk.Combobox(
            fpur,
            textvariable=self._purpose_combo_var,
            values=self.PURPOSE_OPTIONS + [PURPOSE_CUSTOM],
            state="readonly",
            width=52,
        )
        self._purpose_combo.grid(row=0, column=1, padx=4, sticky="ew")
        self._purpose_combo.current(0)
        self._purpose_combo.bind("<<ComboboxSelected>>", self._on_purpose_select)
        self._purpose_custom_frame = ttk.Frame(fpur)
        self._purpose_custom_frame.grid(
            row=1, column=0, columnspan=2, sticky="ew", pady=(4, 0)
        )
        self._purpose_custom_frame.columnconfigure(1, weight=1)
        ttk.Label(
            self._purpose_custom_frame, text="Свой вариант:", anchor="w", width=14
        ).grid(row=0, column=0, sticky="w")
        self._purpose_custom_var = tk.StringVar()
        ttk.Entry(
            self._purpose_custom_frame, textvariable=self._purpose_custom_var, width=46
        ).grid(row=0, column=1, padx=4, sticky="ew")
        self._purpose_custom_frame.grid_remove()

        # ── Установка ─────────────────────────────────
        fi = ttk.LabelFrame(main, text="  Установка  ", padding=8)
        fi.pack(fill="x", **PAD)
        self.install_var = tk.BooleanVar(value=True)
        ttk.Radiobutton(
            fi, text="✓  С установкой", variable=self.install_var, value=True
        ).pack(side="left", padx=16)
        ttk.Radiobutton(
            fi, text="✗  Без установки", variable=self.install_var, value=False
        ).pack(side="left", padx=16)

        # ── Реквизиты ─────────────────────────────────
        fd = ttk.LabelFrame(main, text="  Реквизиты документа  ", padding=8)
        fd.pack(fill="x", **PAD)
        fd.columnconfigure(1, weight=1)
        self._doc = {}
        now = datetime.now()
        for i, (key, label, default) in enumerate(
            [
                ("reg_number", "Рег. номер:", ""),
                ("date_short", "Дата (дд.мм.гггг):", now.strftime("%d.%m.%Y")),
                ("executor", "Исполнитель (журнал):", ""),
                ("output_dir", "Папка сохранения:", os.path.expanduser("~/Desktop")),
            ]
        ):
            ttk.Label(fd, text=label, anchor="w", width=26).grid(
                row=i, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=default)
            self._doc[key] = var
            ttk.Entry(fd, textvariable=var, width=44).grid(
                row=i, column=1, padx=4, sticky="ew"
            )
            if key == "output_dir":
                ttk.Button(fd, text="…", width=3, command=self._pick_outdir).grid(
                    row=i, column=2
                )

        ttk.Button(
            main, text="📄   Создать заявку и добавить в журнал", command=self._generate
        ).pack(pady=10, ipadx=16, ipady=6)

        # Нижний статус (только для общей инфо о загрузке файлов)
        self._load_status_var = tk.StringVar(
            value="Укажите файлы данных и нажмите «Загрузить»"
        )
        self._load_status_lbl = ttk.Label(
            main, textvariable=self._load_status_var, foreground="#555", wraplength=700
        )
        self._load_status_lbl.pack(pady=(0, 12))

    # ── Helpers ───────────────────────────────────────

    def _set_status(self, text, color="gray"):
        clr = {
            "green": "#1a7a1a",
            "red": "#cc0000",
            "orange": "#b06000",
            "gray": "#555",
        }
        self.status_var.set(text)
        self._status_lbl.configure(foreground=clr.get(color, "#555"))

    def _set_load_status(self, text, color="gray"):
        clr = {"green": "#2a7a2a", "red": "#cc0000", "gray": "#555"}
        self._load_status_var.set(text)
        self._load_status_lbl.configure(foreground=clr.get(color, "#555"))

    def _pick_file(self, key, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self._file_vars[key].set(path)
            self.cfg[key] = path
            save_config(self.cfg)

    def _pick_outdir(self):
        d = filedialog.askdirectory()
        if d:
            self._doc["output_dir"].set(d)

    def _open_file(self, path):
        try:
            if sys.platform == "darwin":
                subprocess.run(["open", path])
            elif sys.platform == "win32":
                os.startfile(path)
            else:
                subprocess.run(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть: {e}")

    def _on_purpose_select(self, _=None):
        if self._purpose_combo_var.get() == PURPOSE_CUSTOM:
            self._purpose_custom_frame.grid()
        else:
            self._purpose_custom_frame.grid_remove()

    def _get_purpose(self) -> str:
        sel = self._purpose_combo_var.get()
        if sel == PURPOSE_CUSTOM:
            return self._purpose_custom_var.get().strip()
        return sel

    # ── Load ──────────────────────────────────────────

    def _load_files(self):
        paths = {k: self._file_vars[k].get().strip() for k in self._file_vars}
        errors = []
        loaded = 0

        if paths["phone_book"] and os.path.exists(paths["phone_book"]):
            try:
                self.phone_data = load_phone_book(paths["phone_book"])
                loaded += 1
            except Exception as e:
                errors.append(f"Справочник: {e}")

        if paths["pc_file"] and os.path.exists(paths["pc_file"]):
            try:
                self.pc_data = load_pc_data(paths["pc_file"])
                loaded += 1
            except Exception as e:
                errors.append(f"Актуализация ПК: {e}")

        # Объединяем имена из обоих источников
        self.all_names = sorted(set(self.phone_data.keys()) | set(self.pc_data.keys()))

        if paths["journal"] and os.path.exists(paths["journal"]):
            try:
                self.journal_info = get_journal_info(paths["journal"])
                self._doc["reg_number"].set(self.journal_info["next_reg"])
                if self.journal_info["last_executor"]:
                    self._doc["executor"].set(self.journal_info["last_executor"])
                loaded += 1
            except Exception as e:
                errors.append(f"Журнал: {e}")

        for k in paths:
            self.cfg[k] = paths[k]
        save_config(self.cfg)

        ji = self.journal_info
        if errors:
            self._set_load_status("⚠ " + " | ".join(errors), "red")
        else:
            self._set_load_status(
                f"✓ Загружено {loaded}/3 файлов. "
                f"Сотрудников: {len(self.all_names)}. "
                f"Следующий рег. номер: {ji.get('next_reg','—')}",
                "green",
            )

    # ── Search ────────────────────────────────────────

    def _on_search_change(self, *_):
        q = self.search_var.get().strip().lower()
        self.lb.delete(0, tk.END)
        if len(q) < 2:
            return
        for name in self.all_names:
            if q in name.lower():
                self.lb.insert(tk.END, name)

    def _do_search(self):
        self._on_search_change()

    def _on_select(self, _):
        sel = self.lb.curselection()
        if not sel:
            return
        name = self.lb.get(sel[0])
        self.search_var.set(name)
        self._fill_employee(name)

    def _fill_employee(self, name: str):
        emp = self.phone_data.get(name, {})
        position = emp.get("position", "")
        department = emp.get("department", "")
        abbrev = abbreviate_dept(department)

        self._chief_pos_prefix = chief_position_prefix(emp.get("chief_position", ""))
        self._chief_abbrev = abbrev

        self._emp["emp_name"].set(name)
        self._emp["emp_position_doc"].set(build_position_doc(position, department))
        self._emp["chief_name"].set(emp.get("chief_name", ""))
        self._emp["chief_initials"].set(emp.get("chief_initials", ""))
        self._emp["chief_position_doc"].set(abbrev)

        # Если сотрудник найден только в актуализации — предупреждаем
        if not emp:
            messagebox.showwarning(
                "Сотрудник не в телефонном справочнике",
                f"\u00ab{name}\u00bb найден только в таблице актуализации.\n\n"
                "Заполните вручную следующие поля:\n"
                "  \u2022 Должность и отдел\n"
                "  \u2022 Начальник (ФИО)\n"
                "  \u2022 Инициалы начальника\n"
                "  \u2022 Аббревиатура отдела",
            )

        # Ищем ПК по каноническому имени (без суффиксов "(ноут)" и пр.)
        canonical = _strip_suffix(name)
        pcs = self.pc_data.get(canonical, [])
        # Если не нашли по каноническому — пробуем точное
        if not pcs:
            pcs = self.pc_data.get(name, [])
        self._pcs = pcs

        if not pcs:
            self.pc_combo["values"] = ["Не найден в таблице актуализации"]
            self.pc_combo.current(0)
            self._serial_var.set("")
            self._inv_var.set("")
            self._preview.clear()
            self._cert_var.set("")
            self._set_status("⚠ ПК не найден в таблице актуализации", "orange")
        else:
            labels = [pc["label"] for pc in pcs]
            self.pc_combo["values"] = labels
            self.pc_combo.current(0)
            self._serial_var.set(pcs[0]["serial"])
            self._inv_var.set(pcs[0]["inventory"])
            self._preview.show(pcs)
            self._update_cert_status(pcs)
            if len(pcs) == 1:
                self._set_status(f"✓ Найден: {name}", "green")
            else:
                self._set_status(
                    f"⚠ {len(pcs)} устройства — выберите в «Выбор ПК»", "orange"
                )

    def _update_cert_status(self, pcs: list):
        """Обновляет лейбл статуса сертификата под таблицей превью."""
        if not pcs:
            self._cert_var.set("")
            return
        idx = self.pc_combo.current()
        rec = pcs[idx] if 0 <= idx < len(pcs) else pcs[0]
        date_str = rec["values"][6].strip() if len(rec["values"]) > 6 else ""
        if not date_str or date_str.lower() in ("нет", "-", "none", ""):
            self._cert_var.set("— Сертификат не установлен")
            self._cert_lbl.configure(foreground="#888888")
            return
        try:
            cert_date = datetime.strptime(date_str, "%d.%m.%Y").date()
        except ValueError:
            self._cert_var.set("— Сертификат не установлен")
            self._cert_lbl.configure(foreground="#888888")
            return
        days_left = (cert_date - datetime.now().date()).days
        if days_left < 0:
            self._cert_var.set(
                f"✗ Сертификат истёк {-days_left} дн. назад ({date_str})"
            )
            self._cert_lbl.configure(foreground="#cc0000")
        elif days_left < 60:
            self._cert_var.set(
                f"⚠ Сертификат истекает через {days_left} дн. ({date_str})"
            )
            self._cert_lbl.configure(foreground="#b06000")
        else:
            self._cert_var.set(f"✓ Сертификат действителен до {date_str}")
            self._cert_lbl.configure(foreground="#1a7a1a")

    def _on_pc_select(self, _):
        idx = self.pc_combo.current()
        if 0 <= idx < len(self._pcs):
            self._serial_var.set(self._pcs[idx]["serial"])
            self._inv_var.set(self._pcs[idx]["inventory"])
            self._update_cert_status(self._pcs)

    # ── Generate ──────────────────────────────────────

    def _ask_duplicate(self, path: str, filename: str) -> str:
        win = tk.Toplevel(self)
        win.title("Файл уже существует")
        win.resizable(False, False)
        win.grab_set()
        ttk.Label(
            win,
            text=f"Заявка уже существует:\n{filename}",
            wraplength=400,
            justify="left",
            padding=12,
        ).pack()
        result = tk.StringVar(value="cancel")
        bf = ttk.Frame(win, padding=8)
        bf.pack()

        def choose(v):
            result.set(v)
            win.destroy()

        ttk.Button(
            bf, text="📂  Открыть существующую", command=lambda: choose("open")
        ).grid(row=0, column=0, padx=6, pady=4, sticky="ew")
        ttk.Button(bf, text="🔄  Заменить", command=lambda: choose("replace")).grid(
            row=0, column=1, padx=6, pady=4, sticky="ew"
        )
        ttk.Button(bf, text="➕  Создать ещё одну", command=lambda: choose("new")).grid(
            row=1, column=0, padx=6, pady=4, sticky="ew"
        )
        ttk.Button(bf, text="✖  Отмена", command=lambda: choose("cancel")).grid(
            row=1, column=1, padx=6, pady=4, sticky="ew"
        )
        win.wait_window()
        return result.get()

    def _generate(self):
        emp_name = self._emp["emp_name"].get().strip()
        if not emp_name:
            messagebox.showerror("Ошибка", "Сотрудник не выбран!")
            return
        reg = self._doc["reg_number"].get().strip()
        if not reg:
            messagebox.showerror("Ошибка", "Укажите регистрационный номер!")
            return
        date_str = self._doc["date_short"].get().strip()
        try:
            dt = datetime.strptime(date_str, "%d.%m.%Y")
        except ValueError:
            messagebox.showerror("Ошибка", "Формат даты: дд.мм.гггг")
            return
        purpose = self._get_purpose()
        if not purpose:
            messagebox.showerror("Ошибка", "Укажите цель установки!")
            return

        serial = self._serial_var.get().strip()
        inventory = self._inv_var.get().strip()
        with_install = self.install_var.get()
        if not serial and not inventory:
            if not messagebox.askyesno(
                "Предупреждение",
                "Серийный и инвентарный номера не заполнены.\nПродолжить?",
            ):
                return

        emp_initials = make_initials(emp_name)

        # Проверка дубликата в журнале
        journal_path_check = self._file_vars["journal"].get().strip()
        if journal_path_check and os.path.exists(journal_path_check) and serial:
            dup = check_journal_duplicate(journal_path_check, emp_initials, serial)
            if dup:
                if not messagebox.askyesno(
                    "Дубликат в журнале",
                    f"В журнале уже есть заявка на этого сотрудника с этим ПК:\n{dup}\n\n"
                    f"Продолжить и добавить ещё одну запись?",
                ):
                    return

        abbrev_in_ui = self._emp["chief_position_doc"].get().strip()
        data = {
            "reg_number": reg,
            "date_short": date_str,
            "employee_name": emp_name,
            "employee_position_full": self._emp["emp_position_doc"].get().strip(),
            "chief_name": self._emp["chief_name"].get().strip(),
            "chief_initials": self._emp["chief_initials"].get().strip(),
            "chief_pos_prefix": self._chief_pos_prefix,
            "chief_abbrev": abbrev_in_ui,
            "chief_position_doc": abbrev_in_ui,
            "serial": serial,
            "inventory": inventory,
            "with_install": with_install,
            "purpose": purpose,
            "day": str(dt.day),
            "month": MONTHS_RU[dt.month],
            "year": str(dt.year),
        }

        out_dir = self._doc["output_dir"].get().strip() or os.path.expanduser(
            "~/Desktop"
        )
        os.makedirs(out_dir, exist_ok=True)
        base_name = self._make_filename(emp_initials, serial)
        journal_desc_default = self._journal_description(emp_initials, serial)

        # ── Окно подтверждения ────────────────────────
        confirm = ConfirmWindow(self, base_name, journal_desc_default, out_dir)
        confirmed_name, confirmed_desc = confirm.result
        if confirmed_name is None:
            return  # пользователь нажал Отмена
        base_name = confirmed_name
        output_path = os.path.join(out_dir, base_name)

        if os.path.exists(output_path):
            choice = self._ask_duplicate(output_path, base_name)
            if choice == "open":
                self._open_file(output_path)
                return
            elif choice == "cancel":
                return
            elif choice == "new":
                ts = datetime.now().strftime("%H-%M-%S")
                output_path = os.path.join(
                    out_dir, base_name.replace(".docx", f"_{ts}.docx")
                )

        try:
            self._generate_doc(data, output_path)
        except Exception as e:
            messagebox.showerror("Ошибка при создании документа", str(e))
            return

        journal_path = self._file_vars["journal"].get().strip()
        journal_msg = ""
        if journal_path and os.path.exists(journal_path) and self.journal_info:
            executor_name = self._doc["executor"].get().strip()
            acquired, blocker = acquire_journal_lock(journal_path, executor_name)
            if not acquired:
                if blocker is None:
                    messagebox.showwarning(
                        "Ошибка записи в журнал",
                        "Не удалось заблокировать журнал для записи.\n"
                        "Папка с журналом, возможно, открыта только для чтения.\n"
                        "Обратитесь к администратору.",
                    )
                    journal_msg = (
                        "⚠ Нет прав на запись рядом с журналом — запись не добавлена."
                    )
                else:
                    messagebox.showwarning(
                        "Журнал заблокирован",
                        f"Журнал сейчас редактирует другой пользователь:\n{blocker}\n\n"
                        f"Подождите и попробуйте снова.",
                    )
                    journal_msg = (
                        f"⚠ Журнал заблокирован: {blocker} — запись не добавлена."
                    )
            else:
                try:
                    entry = {
                        "pp": self.journal_info["next_pp"],
                        "date": date_str,
                        "reg": reg,
                        "description": confirmed_desc,
                        "executor": self._doc["executor"].get().strip(),
                        "note": "Акт установки" if with_install else "",
                    }
                    journal_msg = _write_journal_with_retry(
                        journal_path, entry, self.journal_info["last_row_idx"]
                    )
                    if journal_msg.startswith("✓"):
                        self.journal_info = get_journal_info(journal_path)
                        self._doc["reg_number"].set(self.journal_info["next_reg"])
                finally:
                    release_journal_lock(journal_path)
        else:
            journal_msg = "Журнал не указан — запись не добавлена."

        self._set_status(f"✓ Создан: {os.path.basename(output_path)}", "green")
        if messagebox.askyesno(
            "Готово!",
            f"Заявка создана:\n{output_path}\n\n{journal_msg}\n\nОткрыть файл?",
        ):
            self._open_file(output_path)


# ══════════════════════════════════════════════════════
#  ConfirmWindow — окно подтверждения перед генерацией
# ══════════════════════════════════════════════════════


class ConfirmWindow(tk.Toplevel):
    """Показывает имя файла и описание журнала перед генерацией.
    Пользователь может отредактировать оба поля.
    Возвращает (filename, journal_desc) или (None, None) при отмене."""

    def __init__(self, parent, filename: str, journal_desc: str, out_dir: str):
        super().__init__(parent)
        self.title("Подтверждение")
        self.resizable(True, False)
        self.grab_set()

        self._result_filename   = None
        self._result_journal    = None

        PAD = dict(padx=10, pady=4)

        ttk.Label(
            self,
            text="Проверьте и при необходимости отредактируйте имя файла\nи описание для журнала регистрации.",
            justify="left",
        ).pack(anchor="w", padx=12, pady=(12, 4))

        # ── Имя файла ─────────────────────────────────
        ff = ttk.LabelFrame(self, text="  Имя файла  ", padding=8)
        ff.pack(fill="x", **PAD)
        ff.columnconfigure(0, weight=1)
        self._fname_var = tk.StringVar(value=filename)
        ttk.Entry(ff, textvariable=self._fname_var, width=72).grid(
            row=0, column=0, sticky="ew"
        )
        ttk.Label(
            ff,
            text=f"Папка: {out_dir}",
            foreground="#666",
            font=("TkDefaultFont", 9),
        ).grid(row=1, column=0, sticky="w", pady=(4, 0))

        # ── Описание журнала ──────────────────────────
        fj = ttk.LabelFrame(self, text="  Запись в журнал регистрации  ", padding=8)
        fj.pack(fill="x", **PAD)
        fj.columnconfigure(0, weight=1)
        self._jdesc_var = tk.StringVar(value=journal_desc)
        ttk.Entry(fj, textvariable=self._jdesc_var, width=72).grid(
            row=0, column=0, sticky="ew"
        )

        # ── Кнопки ───────────────────────────────────
        bf = ttk.Frame(self)
        bf.pack(pady=(8, 12))
        ttk.Button(
            bf, text="✓  Создать", width=18, command=self._ok
        ).pack(side="left", padx=8)
        ttk.Button(
            bf, text="✖  Отмена", width=14, command=self._cancel
        ).pack(side="left", padx=8)

        self.bind("<Return>", lambda _: self._ok())
        self.bind("<Escape>", lambda _: self._cancel())

        # Центрируем относительно родителя
        self.update_idletasks()
        pw, ph = parent.winfo_width(), parent.winfo_height()
        px, py = parent.winfo_rootx(), parent.winfo_rooty()
        ww, wh = self.winfo_width(), self.winfo_height()
        self.geometry(f"+{px + (pw - ww)//2}+{py + (ph - wh)//2}")

        self.wait_window(self)

    def _ok(self):
        fname = self._fname_var.get().strip()
        if not fname.endswith(".docx"):
            fname += ".docx"
        self._result_filename = fname
        self._result_journal  = self._jdesc_var.get().strip()
        self.destroy()

    def _cancel(self):
        self.destroy()

    @property
    def result(self):
        """Возвращает (filename, journal_desc) или (None, None)."""
        return self._result_filename, self._result_journal


# ══════════════════════════════════════════════════════
#  PKIApp — заявки ViPNet PKI Client
# ══════════════════════════════════════════════════════


class PKIApp(BaseZayavkaApp):
    PURPOSE_OPTIONS = ["для подписания ЭП на портале ЕЦП"]

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Генератор заявок ViPNet PKI Client")

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        return (
            f"Заявка на обучение СКЗИ  ViPNet CSP в составе "
            f"ПО ViPNet PKI Client {emp_initials} ({serial}).docx"
        )

    def _generate_doc(self, data: dict, output_path: str):
        generate_zayavka(data, output_path)

    def _journal_description(self, emp_initials: str, serial: str) -> str:
        return f"Заявка на обучение PKI Client {emp_initials} ({serial})"


# ══════════════════════════════════════════════════════
#  CSPApp — заявки ViPNet CSP Client
# ══════════════════════════════════════════════════════


class CSPApp(BaseZayavkaApp):
    PURPOSE_OPTIONS = [
        "работой с ПК: ПТК КС, Элардо, АРМ БПИ, портал Казначейства, "
        "для взаимодействия по защищенному каналу по средствам VipNet «Деловая почта»"
    ]

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Генератор заявок ViPNet CSP Client")

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        return (
            f"Заявка на установку СКЗИ  ViPNet CSP в составе "
            f"СПО ViPNet Client {emp_initials} ({serial}).docx"
        )

    def _generate_doc(self, data: dict, output_path: str):
        generate_csp_zayavka(data, output_path)

    def _journal_description(self, emp_initials: str, serial: str) -> str:
        return f"Акт установки ViPNet CSP в составе СПО VipNet Client {emp_initials} ({serial})"


# ══════════════════════════════════════════════════════
#  CSPPtkApp — заявки ViPNet CSP (ПТК КС)
# ══════════════════════════════════════════════════════


class CSPPtkApp(BaseZayavkaApp):
    PURPOSE_OPTIONS = ["Работой в ПТК КС"]

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Генератор заявок ViPNet CSP (ПТК КС)")

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        return f"Заявка на установку СКЗИ ViPNet CSP " f"{emp_initials} ({serial}).docx"

    def _generate_doc(self, data: dict, output_path: str):
        generate_ptk_zayavka(data, output_path)

    def _journal_description(self, emp_initials: str, serial: str) -> str:
        return f"Акт установки СКЗИ VipNet CSP {emp_initials} ({serial})"


# ══════════════════════════════════════════════════════
#  KriptoproApp — заявки КриптоПРО CSP
# ══════════════════════════════════════════════════════


class KriptoproApp(BaseZayavkaApp):
    PURPOSE_OPTIONS = ["работой в ЕЦП, работа на портале Казначейства"]

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Генератор заявок КриптоПРО CSP")

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        return f"Заявка на обучение СКЗИ КриптоПРО CSP {emp_initials} ({serial}).docx"

    def _generate_doc(self, data: dict, output_path: str):
        generate_kriptopro_zayavka(data, output_path)

    def _journal_description(self, emp_initials: str, serial: str) -> str:
        return f"Заявка на обучение СКЗИ КриптоПРО CSP {emp_initials} ({serial})"


# ══════════════════════════════════════════════════════
#  ECPApp — Заявки на доступ в ГИС ЕЦП
# ══════════════════════════════════════════════════════


class ECPApp(tk.Toplevel):

    def __init__(self, launcher):
        super().__init__(launcher)
        self._launcher = launcher
        self.title("Заявка на доступ в ГИС ЕЦП")
        self.geometry("800x920")
        self.resizable(True, True)
        self.cfg = load_config()
        self.phone_data = {}
        self.journal_info = {}
        self.all_names = []
        self._chief_pos_prefix = ""
        self._chief_abbrev = ""
        self.protocol("WM_DELETE_WINDOW", self._back)
        self._build_ui()
        self._load_files()
        self._bind_clipboard()

    def _back(self):
        self._launcher.deiconify()
        self.destroy()

    def _bind_clipboard(self):
        pass  # clipboard handled at root level in LauncherWindow

    # ── UI ────────────────────────────────────────────

    def _build_ui(self):
        main_canvas = tk.Canvas(
            self, borderwidth=0, highlightthickness=0, yscrollincrement=20
        )
        vbar = ttk.Scrollbar(self, orient="vertical", command=main_canvas.yview)
        main_canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        main_canvas.pack(side="left", fill="both", expand=True)

        main = tk.Frame(main_canvas)
        self._main_win = main_canvas.create_window((0, 0), window=main, anchor="nw")
        main.bind(
            "<Configure>",
            lambda _: main_canvas.configure(scrollregion=main_canvas.bbox("all")),
        )
        main_canvas.bind(
            "<Configure>",
            lambda e: main_canvas.itemconfig(self._main_win, width=e.width),
        )

        def _sd(event):
            return int(-1 * (event.delta / 120)) if abs(event.delta) >= 120 else (-1 if event.delta > 0 else 1)

        def _on_mousewheel(event):
            if not event.delta or not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(_sd(event), "units")

        def _on_wheel_linux(event, direction):
            if not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(direction, "units")

        self.bind_all("<MouseWheel>", _on_mousewheel)
        self.bind_all("<Button-4>", lambda e: _on_wheel_linux(e, -1))
        self.bind_all("<Button-5>", lambda e: _on_wheel_linux(e, 1))

        PAD = dict(padx=8, pady=4)

        ttk.Button(main, text="← Назад", command=self._back).pack(
            anchor="w", padx=8, pady=(6, 0)
        )

        # ── Файлы ─────────────────────────────────────
        ff = ttk.LabelFrame(main, text="  Файлы данных  ", padding=8)
        ff.pack(fill="x", **PAD)
        file_specs = [
            ("phone_book", "Телефонный справочник (.xls):", [("XLS", "*.xls"), ("Все", "*.*")]),
            ("journal",    "Журнал регистрации (.docx):",   [("DOCX", "*.docx"), ("Все", "*.*")]),
        ]
        self._file_vars = {}
        for ri, (key, label, ftypes) in enumerate(file_specs):
            ttk.Label(ff, text=label, width=30, anchor="w").grid(row=ri, column=0, sticky="w", pady=2)
            var = tk.StringVar(value=self.cfg.get(key, ""))
            self._file_vars[key] = var
            ttk.Entry(ff, textvariable=var, width=42).grid(row=ri, column=1, padx=4)
            ttk.Button(
                ff, text="…", width=3,
                command=lambda k=key, ft=ftypes: self._pick_file(k, ft),
            ).grid(row=ri, column=2)
        ttk.Button(
            ff, text="↺  Загрузить / обновить файлы", command=self._load_files
        ).grid(row=2, column=0, columnspan=3, pady=(6, 0))

        # ── Поиск ─────────────────────────────────────
        fs = ttk.LabelFrame(main, text="  Поиск сотрудника  ", padding=8)
        fs.pack(fill="x", **PAD)
        fs.columnconfigure(1, weight=1)
        ttk.Label(fs, text="ФИО:").grid(row=0, column=0, sticky="w")
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self._on_search_change)
        ttk.Entry(fs, textvariable=self.search_var, width=40).grid(row=0, column=1, padx=4, sticky="ew")
        ttk.Button(fs, text="Найти", command=self._do_search).grid(row=0, column=2)
        self.status_var = tk.StringVar(value="")
        self._status_lbl = tk.Label(
            fs, textvariable=self.status_var, foreground="#555",
            wraplength=550, anchor="w", justify="left",
            font=("TkDefaultFont", 11, "bold"),
        )
        self._status_lbl.grid(row=0, column=3, padx=(12, 4), sticky="ew")
        fs.columnconfigure(3, weight=1)
        self.lb = tk.Listbox(fs, height=10, width=50, font=("TkDefaultFont", 10))
        self.lb.grid(row=1, column=0, columnspan=2, pady=(4, 0), sticky="ew")
        self.lb.bind("<<ListboxSelect>>", self._on_select)
        sb = ttk.Scrollbar(fs, orient="vertical", command=self.lb.yview)
        sb.grid(row=1, column=2, sticky="ns", pady=(4, 0))
        self.lb.configure(yscrollcommand=sb.set)
        self.lb.bind("<MouseWheel>", lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(int(-1*(e.delta/120)) if abs(e.delta)>=120 else (-1 if e.delta>0 else 1), "units"), "break")[1] if e.delta else "break")
        self.lb.bind("<Button-4>", lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(-1, "units"), "break")[1])
        self.lb.bind("<Button-5>", lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(1, "units"), "break")[1])

        # ── Данные сотрудника ─────────────────────────
        fe = ttk.LabelFrame(main, text="  Данные сотрудника (для документа)  ", padding=8)
        fe.pack(fill="x", **PAD)
        fe.columnconfigure(1, weight=1)
        self._emp = {}
        missing_frame = ttk.Frame(fe)
        missing_frame.grid(row=0, column=0, columnspan=2, sticky="ew")
        self._missing_var = tk.StringVar(value="")
        self._missing_lbl = tk.Label(
            missing_frame, textvariable=self._missing_var,
            foreground="#cc0000", anchor="w", justify="left",
            font=("TkDefaultFont", 10), wraplength=600,
        )
        self._missing_lbl.pack(fill="x")

        fields = [
            ("emp_name",      "ФИО сотрудника:"),
            ("emp_dept",      "Управление / Отдел:"),
            ("emp_position",  "Должность:"),
            ("chief_initials","Инициалы начальника (подпись):"),
            ("phone_short",   "Телефон (4 цифры):"),
        ]
        for i, (key, label) in enumerate(fields, start=1):
            ttk.Label(fe, text=label, anchor="w", width=34).grid(row=i, column=0, sticky="w", pady=2)
            var = tk.StringVar()
            self._emp[key] = var
            ttk.Entry(fe, textvariable=var, width=46).grid(row=i, column=1, padx=4, sticky="ew")

        # ── Среда ─────────────────────────────────────
        fenv = ttk.LabelFrame(main, text="  Среда  ", padding=8)
        fenv.pack(fill="x", **PAD)
        self.env_var = tk.StringVar(value="prod")
        ttk.Radiobutton(fenv, text="ГИС ЕЦП (ИС КНД) Прод", variable=self.env_var, value="prod").pack(side="left", padx=16)
        ttk.Radiobutton(fenv, text="ГИС ЕЦП ТЕСТ",           variable=self.env_var, value="test").pack(side="left", padx=16)

        # ── Действия ──────────────────────────────────
        fact = ttk.LabelFrame(main, text="  Действие  ", padding=8)
        fact.pack(fill="x", **PAD)
        self.action_unlock_var = tk.BooleanVar(value=False)
        self.action_change_pass_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(fact, text="Разблокировать пользователя", variable=self.action_unlock_var).pack(anchor="w", padx=16)
        ttk.Checkbutton(fact, text="Сменить пароль",              variable=self.action_change_pass_var).pack(anchor="w", padx=16)

        # ── Реквизиты ─────────────────────────────────
        fd = ttk.LabelFrame(main, text="  Реквизиты документа  ", padding=8)
        fd.pack(fill="x", **PAD)
        fd.columnconfigure(1, weight=1)
        self._doc = {}
        now = datetime.now()
        for i, (key, label, default) in enumerate([
            ("reg_number", "Рег. номер:",         ""),
            ("date_short", "Дата (дд.мм.гггг):",  now.strftime("%d.%m.%Y")),
            ("executor",   "Исполнитель (журнал):", ""),
            ("output_dir", "Папка сохранения:",    os.path.expanduser("~/Desktop")),
        ]):
            ttk.Label(fd, text=label, anchor="w", width=26).grid(row=i, column=0, sticky="w", pady=2)
            var = tk.StringVar(value=default)
            self._doc[key] = var
            ttk.Entry(fd, textvariable=var, width=44).grid(row=i, column=1, padx=4, sticky="ew")
            if key == "output_dir":
                ttk.Button(fd, text="…", width=3, command=self._pick_outdir).grid(row=i, column=2)

        ttk.Button(
            main, text="📄   Создать заявку и добавить в журнал", command=self._generate
        ).pack(pady=10, ipadx=16, ipady=6)

        self._load_status_var = tk.StringVar(value="Укажите файлы данных и нажмите «Загрузить»")
        self._load_status_lbl = ttk.Label(
            main, textvariable=self._load_status_var, foreground="#555", wraplength=700
        )
        self._load_status_lbl.pack(pady=(0, 12))

    # ── Helpers ───────────────────────────────────────

    def _set_status(self, text, color="gray"):
        clr = {"green": "#1a7a1a", "red": "#cc0000", "orange": "#b06000", "gray": "#555"}
        self.status_var.set(text)
        self._status_lbl.configure(foreground=clr.get(color, "#555"))

    def _set_load_status(self, text, color="gray"):
        clr = {"green": "#2a7a2a", "red": "#cc0000", "gray": "#555"}
        self._load_status_var.set(text)
        self._load_status_lbl.configure(foreground=clr.get(color, "#555"))

    def _pick_file(self, key, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self._file_vars[key].set(path)
            self.cfg[key] = path
            save_config(self.cfg)

    def _pick_outdir(self):
        d = filedialog.askdirectory()
        if d:
            self._doc["output_dir"].set(d)

    def _open_file(self, path):
        try:
            if sys.platform == "darwin":
                subprocess.run(["open", path])
            elif sys.platform == "win32":
                os.startfile(path)
            else:
                subprocess.run(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть: {e}")

    # ── Load ──────────────────────────────────────────

    def _load_files(self):
        paths = {k: self._file_vars[k].get().strip() for k in self._file_vars}
        errors = []
        loaded = 0

        if paths["phone_book"] and os.path.exists(paths["phone_book"]):
            try:
                self.phone_data = load_phone_book(paths["phone_book"])
                self.all_names = sorted(self.phone_data.keys())
                loaded += 1
            except Exception as e:
                errors.append(f"Справочник: {e}")

        if paths["journal"] and os.path.exists(paths["journal"]):
            try:
                self.journal_info = get_journal_info(paths["journal"])
                self._doc["reg_number"].set(self.journal_info["next_reg"])
                if self.journal_info["last_executor"]:
                    self._doc["executor"].set(self.journal_info["last_executor"])
                loaded += 1
            except Exception as e:
                errors.append(f"Журнал: {e}")

        for k in paths:
            self.cfg[k] = paths[k]
        save_config(self.cfg)

        if errors:
            self._set_load_status("⚠ " + " | ".join(errors), "red")
        else:
            ji = self.journal_info
            self._set_load_status(
                f"✓ Загружено {loaded}/2 файлов. "
                f"Сотрудников: {len(self.all_names)}. "
                f"Следующий рег. номер: {ji.get('next_reg', '—')}",
                "green",
            )

    # ── Search ────────────────────────────────────────

    def _on_search_change(self, *_):
        q = self.search_var.get().strip().lower()
        self.lb.delete(0, tk.END)
        if len(q) < 2:
            return
        for name in self.all_names:
            if q in name.lower():
                self.lb.insert(tk.END, name)

    def _do_search(self):
        self._on_search_change()

    def _on_select(self, _):
        sel = self.lb.curselection()
        if not sel:
            return
        name = self.lb.get(sel[0])
        self.search_var.set(name)
        self._fill_employee(name)

    def _fill_employee(self, name: str):
        emp = self.phone_data.get(name, {})
        position  = emp.get("position", "")
        dept      = emp.get("department", "")
        phone_raw = emp.get("phone", "")
        abbrev    = abbreviate_dept(dept)

        self._chief_pos_prefix = chief_position_prefix(emp.get("chief_position", ""))
        self._chief_abbrev = abbrev

        # Заполняем поля
        self._emp["emp_name"].set(name)
        self._emp["emp_dept"].set(dept)
        self._emp["emp_position"].set(position)
        self._emp["chief_initials"].set(emp.get("chief_initials", ""))
        self._emp["phone_short"].set(_extract_phone_short(phone_raw))

        # Проверяем чего не хватает
        missing = []
        if not dept:       missing.append("Управление / Отдел")
        if not position:   missing.append("Должность")
        if not emp.get("chief_initials"): missing.append("Инициалы начальника")
        if not phone_raw:  missing.append("Телефон")

        if missing:
            self._missing_var.set("⚠ Не найдено в справочнике: " + ", ".join(missing) + " — заполните вручную")
        else:
            self._missing_var.set("")

        if not emp:
            self._set_status(f"⚠ Сотрудник не в справочнике", "orange")
        else:
            self._set_status(f"✓ Найден: {name}", "green")

    # ── Generate ──────────────────────────────────────

    def _ask_duplicate(self, path: str, filename: str) -> str:
        win = tk.Toplevel(self)
        win.title("Файл уже существует")
        win.resizable(False, False)
        win.grab_set()
        ttk.Label(win, text=f"Заявка уже существует:\n{filename}", wraplength=400, justify="left", padding=12).pack()
        result = tk.StringVar(value="cancel")
        bf = ttk.Frame(win, padding=8)
        bf.pack()
        def choose(v): result.set(v); win.destroy()
        ttk.Button(bf, text="📂  Открыть существующую", command=lambda: choose("open")).grid(row=0, column=0, padx=6, pady=4, sticky="ew")
        ttk.Button(bf, text="🔄  Заменить",             command=lambda: choose("replace")).grid(row=0, column=1, padx=6, pady=4, sticky="ew")
        ttk.Button(bf, text="➕  Создать ещё одну",      command=lambda: choose("new")).grid(row=1, column=0, padx=6, pady=4, sticky="ew")
        ttk.Button(bf, text="✖  Отмена",                command=lambda: choose("cancel")).grid(row=1, column=1, padx=6, pady=4, sticky="ew")
        win.wait_window()
        return result.get()

    def _generate(self):
        emp_name = self._emp["emp_name"].get().strip()
        if not emp_name:
            messagebox.showerror("Ошибка", "Сотрудник не выбран!")
            return
        reg = self._doc["reg_number"].get().strip()
        if not reg:
            messagebox.showerror("Ошибка", "Укажите регистрационный номер!")
            return
        date_str = self._doc["date_short"].get().strip()
        try:
            datetime.strptime(date_str, "%d.%m.%Y")
        except ValueError:
            messagebox.showerror("Ошибка", "Формат даты: дд.мм.гггг")
            return

        emp_initials = make_initials(emp_name)

        data = {
            "reg_number":     reg,
            "date_short":     date_str,
            "emp_name":       emp_name,
            "emp_dept":       self._emp["emp_dept"].get().strip(),
            "emp_position":   self._emp["emp_position"].get().strip(),
            "chief_initials": self._emp["chief_initials"].get().strip(),
            "chief_pos_prefix": self._chief_pos_prefix,
            "chief_abbrev":   self._chief_abbrev,
            "phone_short":    self._emp["phone_short"].get().strip(),
            "env":            self.env_var.get(),
            "action_unlock":      self.action_unlock_var.get(),
            "action_change_pass": self.action_change_pass_var.get(),
        }

        out_dir = self._doc["output_dir"].get().strip() or os.path.expanduser("~/Desktop")
        os.makedirs(out_dir, exist_ok=True)
        base_name = f"Заявка ЕЦП {emp_initials}.docx"
        journal_desc_default = f"Заявка ЕЦП {emp_initials}"

        confirm = ConfirmWindow(self, base_name, journal_desc_default, out_dir)
        confirmed_name, confirmed_desc = confirm.result
        if confirmed_name is None:
            return
        base_name = confirmed_name
        output_path = os.path.join(out_dir, base_name)

        if os.path.exists(output_path):
            choice = self._ask_duplicate(output_path, base_name)
            if choice == "open":
                self._open_file(output_path)
                return
            elif choice == "cancel":
                return
            elif choice == "new":
                ts = datetime.now().strftime("%H-%M-%S")
                output_path = os.path.join(out_dir, base_name.replace(".docx", f"_{ts}.docx"))

        try:
            generate_ecp_zayavka(data, output_path)
        except Exception as e:
            messagebox.showerror("Ошибка при создании документа", str(e))
            return

        journal_path = self._file_vars["journal"].get().strip()
        journal_msg = ""
        if journal_path and os.path.exists(journal_path) and self.journal_info:
            executor_name = self._doc["executor"].get().strip()
            acquired, blocker = acquire_journal_lock(journal_path, executor_name)
            if not acquired:
                if blocker is None:
                    messagebox.showwarning("Ошибка записи в журнал",
                        "Не удалось заблокировать журнал для записи.\n"
                        "Папка с журналом, возможно, открыта только для чтения.\n"
                        "Обратитесь к администратору.")
                    journal_msg = "⚠ Нет прав на запись рядом с журналом — запись не добавлена."
                else:
                    messagebox.showwarning("Журнал заблокирован",
                        f"Журнал сейчас редактирует другой пользователь:\n{blocker}\n\nПодождите и попробуйте снова.")
                    journal_msg = f"⚠ Журнал заблокирован: {blocker} — запись не добавлена."
            else:
                try:
                    entry = {
                        "pp":          self.journal_info["next_pp"],
                        "date":        date_str,
                        "reg":         reg,
                        "description": confirmed_desc,
                        "executor":    executor_name,
                        "note":        "",
                    }
                    journal_msg = _write_journal_with_retry(
                        journal_path, entry, self.journal_info["last_row_idx"]
                    )
                    if journal_msg.startswith("✓"):
                        self.journal_info = get_journal_info(journal_path)
                        self._doc["reg_number"].set(self.journal_info["next_reg"])
                finally:
                    release_journal_lock(journal_path)
        else:
            journal_msg = "Журнал не указан — запись не добавлена."

        self._set_status(f"✓ Создан: {os.path.basename(output_path)}", "green")
        if messagebox.askyesno("Готово!",
            f"Заявка создана:\n{output_path}\n\n{journal_msg}\n\nОткрыть файл?"):
            self._open_file(output_path)


# ══════════════════════════════════════════════════════
#  AktPKIApp — Акт установки СКЗИ ViPNet PKI Client
# ══════════════════════════════════════════════════════


class AktPKIApp(tk.Toplevel):

    OOTIKZI_NAME = "Шеляпин Константин Александрович"
    OOTIKZI_POSITION = (
        "главный специалист-эксперт отдела организационно-технической "
        "и криптографической защиты информации"
    )
    OOTIKZI_INITIALS = "Шеляпин К.А."

    def __init__(self, launcher):
        super().__init__(launcher)
        self._launcher = launcher
        self.title("Акт установки СКЗИ ViPNet PKI Client")
        self.geometry("800x980")
        self.resizable(True, True)
        self.cfg = load_config()
        self.phone_data = {}
        self.pc_data = {}
        self.journal_info = {}
        self.all_names = []
        self._pcs = []
        self.protocol("WM_DELETE_WINDOW", self._back)
        self._build_ui()
        self._load_files()
        self._bind_clipboard()

    def _back(self):
        self._launcher.deiconify()
        self.destroy()

    def _bind_clipboard(self):
        pass  # clipboard handled at root level in LauncherWindow

    # ── UI ────────────────────────────────────────────

    def _build_ui(self):
        main_canvas = tk.Canvas(
            self, borderwidth=0, highlightthickness=0, yscrollincrement=20
        )
        vbar = ttk.Scrollbar(self, orient="vertical", command=main_canvas.yview)
        main_canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        main_canvas.pack(side="left", fill="both", expand=True)

        main = tk.Frame(main_canvas)
        self._main_win = main_canvas.create_window((0, 0), window=main, anchor="nw")
        main.bind(
            "<Configure>",
            lambda _: main_canvas.configure(scrollregion=main_canvas.bbox("all")),
        )
        main_canvas.bind(
            "<Configure>",
            lambda e: main_canvas.itemconfig(self._main_win, width=e.width),
        )

        # Скролл колёсиком — всегда главная страница
        def _sd(event):
            return int(-1 * (event.delta / 120)) if abs(event.delta) >= 120 else (-1 if event.delta > 0 else 1)

        def _on_mousewheel(event):
            if not event.delta:
                return
            if not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(_sd(event), 'units')

        def _on_wheel_linux(event, direction):
            if not main_canvas.winfo_exists():
                return
            main_canvas.yview_scroll(direction, 'units')

        self.bind_all('<MouseWheel>', _on_mousewheel)
        self.bind_all('<Button-4>', lambda e: _on_wheel_linux(e, -1))
        self.bind_all('<Button-5>', lambda e: _on_wheel_linux(e,  1))

        PAD = dict(padx=8, pady=4)

        ttk.Button(main, text="← Назад", command=self._back).pack(
            anchor="w", padx=8, pady=(6, 0)
        )

        # ── Файлы ─────────────────────────────────────
        ff = ttk.LabelFrame(main, text="  Файлы данных  ", padding=8)
        ff.pack(fill="x", **PAD)
        file_specs = [
            (
                "phone_book",
                "Телефонный справочник (.xls):",
                [("XLS", "*.xls"), ("Все", "*.*")],
            ),
            (
                "pc_file",
                "Актуализация ПК (.xlsx):",
                [("XLSX", "*.xlsx"), ("Все", "*.*")],
            ),
            (
                "journal",
                "Журнал регистрации (.docx):",
                [("DOCX", "*.docx"), ("Все", "*.*")],
            ),
        ]
        self._file_vars = {}
        for ri, (key, label, ftypes) in enumerate(file_specs):
            ttk.Label(ff, text=label, width=34, anchor="w").grid(
                row=ri, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=self.cfg.get(key, ""))
            self._file_vars[key] = var
            ttk.Entry(ff, textvariable=var, width=38).grid(row=ri, column=1, padx=4)
            ttk.Button(
                ff,
                text="…",
                width=3,
                command=lambda k=key, ft=ftypes: self._pick_file(k, ft),
            ).grid(row=ri, column=2)
        ttk.Button(
            ff, text="↺  Загрузить / обновить файлы", command=self._load_files
        ).grid(row=4, column=0, columnspan=3, pady=(6, 0))

        # ── Поиск сотрудника ──────────────────────────
        fs = ttk.LabelFrame(main, text="  Поиск сотрудника  ", padding=8)
        fs.pack(fill="x", **PAD)
        fs.columnconfigure(1, weight=1)
        ttk.Label(fs, text="ФИО:").grid(row=0, column=0, sticky="w")
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self._on_search_change)
        ttk.Entry(fs, textvariable=self.search_var, width=40).grid(
            row=0, column=1, padx=4, sticky="ew"
        )
        ttk.Button(fs, text="Найти", command=self._on_search_change).grid(
            row=0, column=2
        )
        self.status_var = tk.StringVar()
        self._status_lbl = tk.Label(
            fs,
            textvariable=self.status_var,
            foreground="#555",
            wraplength=500,
            anchor="w",
            font=("TkDefaultFont", 11, "bold"),
        )
        self._status_lbl.grid(row=0, column=3, padx=(12, 4), sticky="ew")
        fs.columnconfigure(3, weight=1)
        self.lb = tk.Listbox(fs, height=10, width=50, font=("TkDefaultFont", 10))
        self.lb.grid(row=1, column=0, columnspan=2, pady=(4, 0), sticky="ew")
        self.lb.bind("<<ListboxSelect>>", self._on_select)
        sb = ttk.Scrollbar(fs, orient="vertical", command=self.lb.yview)
        sb.grid(row=1, column=2, sticky="ns", pady=(4, 0))
        self.lb.configure(yscrollcommand=sb.set)
        self.lb.bind('<MouseWheel>', lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(int(-1*(e.delta/120)) if abs(e.delta)>=120 else (-1 if e.delta>0 else 1), 'units'), 'break')[1] if e.delta else 'break')
        self.lb.bind('<Button-4>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(-1, 'units'), 'break')[1])
        self.lb.bind('<Button-5>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll( 1, 'units'), 'break')[1])

        # Превью таблица Excel
        self._preview = ExcelPreviewTable(main)
        self._preview.pack(fill="x", padx=8, pady=(0, 4))

        # ── Данные ПК ─────────────────────────────────
        fpc = ttk.LabelFrame(main, text="  Компьютер  ", padding=8)
        fpc.pack(fill="x", **PAD)
        fpc.columnconfigure(1, weight=1)
        ttk.Label(fpc, text="Выбор ПК:", anchor="w", width=22).grid(
            row=0, column=0, sticky="w", pady=2
        )
        self.pc_combo = ttk.Combobox(fpc, width=50, state="readonly")
        self.pc_combo.grid(row=0, column=1, padx=4, sticky="ew")
        self.pc_combo.bind("<<ComboboxSelected>>", self._on_pc_select)
        self._serial_var = tk.StringVar()
        self._inv_var = tk.StringVar()
        self._floor_var = tk.StringVar()
        self._room_var = tk.StringVar()
        for ri, (txt, var) in enumerate(
            [
                ("Серийный номер:", self._serial_var),
                ("Инвентарный номер:", self._inv_var),
                ("Этаж:", self._floor_var),
                ("Кабинет:", self._room_var),
            ],
            1,
        ):
            ttk.Label(fpc, text=txt, anchor="w", width=22).grid(
                row=ri, column=0, sticky="w", pady=2
            )
            ttk.Entry(fpc, textvariable=var, width=40).grid(
                row=ri, column=1, padx=4, sticky="ew"
            )

        # ── Поиск заявки в журнале ────────────────────
        fj = ttk.LabelFrame(main, text="  Ссылка на заявку (из журнала)  ", padding=8)
        fj.pack(fill="x", **PAD)
        fj.columnconfigure(1, weight=1)
        ttk.Label(fj, text="Поиск по ФИО:", anchor="w").grid(
            row=0, column=0, sticky="w", pady=2
        )
        self._jsearch_var = tk.StringVar()
        ttk.Entry(fj, textvariable=self._jsearch_var, width=60).grid(
            row=0, column=1, columnspan=2, padx=4, sticky="ew"
        )
        # Live-поиск: результаты появляются сразу при вводе
        self._jsearch_debounce = None

        def _jsearch_debounced(*_):
            if self._jsearch_debounce:
                self.after_cancel(self._jsearch_debounce)
            self._jsearch_debounce = self.after(300, self._on_journal_search)

        self._jsearch_var.trace_add("write", _jsearch_debounced)
        self._jstatus_var = tk.StringVar()
        tk.Label(
            fj,
            textvariable=self._jstatus_var,
            foreground="#b06000",
            anchor="w",
            font=("TkDefaultFont", 10),
        ).grid(row=1, column=0, columnspan=3, sticky="w")
        # Таблица результатов
        jcols = ("pp", "date", "reg", "desc", "executor")
        self._jtree = ttk.Treeview(fj, columns=jcols, show="headings", height=18)
        self._jtree.heading("pp", text="№")
        self._jtree.heading("date", text="Дата")
        self._jtree.heading("reg", text="Рег. номер")
        self._jtree.heading("desc", text="Наименование документа")
        self._jtree.heading("executor", text="Исполнитель")
        self._jtree.column("pp", width=40, stretch=False)
        self._jtree.column("date", width=90, stretch=False)
        self._jtree.column("reg", width=130, stretch=False)
        self._jtree.column("desc", width=320, stretch=True)
        self._jtree.column("executor", width=110, stretch=False)
        jsb = ttk.Scrollbar(fj, orient="vertical", command=self._jtree.yview)
        self._jtree.configure(yscrollcommand=jsb.set)
        self._jtree.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=(4, 0))
        jsb.grid(row=2, column=2, sticky="ns", pady=(4, 0))
        fj.rowconfigure(2, weight=1)
        self._jtree.bind("<<TreeviewSelect>>", self._on_jresult_select)
        self._jtree.bind('<MouseWheel>', lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(int(-1*(e.delta/120)) if abs(e.delta)>=120 else (-1 if e.delta>0 else 1), 'units'), 'break')[1] if e.delta else 'break')
        self._jtree.bind('<Button-4>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll(-1, 'units'), 'break')[1])
        self._jtree.bind('<Button-5>',   lambda e: (main_canvas.winfo_exists() and main_canvas.yview_scroll( 1, 'units'), 'break')[1])
        ttk.Label(fj, text="Ссылка на заявку:", anchor="w").grid(
            row=3, column=0, sticky="w", pady=2
        )
        self._zayavka_ref_var = tk.StringVar()
        ttk.Entry(fj, textvariable=self._zayavka_ref_var, width=50).grid(
            row=3, column=1, columnspan=2, padx=4, sticky="ew"
        )

        # ── ООТиКЗИ ───────────────────────────────────
        fo = ttk.LabelFrame(
            main, text="  Сотрудник ООТ и КЗИ (составитель Акта)  ", padding=8
        )
        fo.pack(fill="x", **PAD)
        fo.columnconfigure(1, weight=1)
        self._ootikzi = {}
        for i, (key, label, default) in enumerate(
            [
                ("name", "ФИО:", self.OOTIKZI_NAME),
                ("position", "Должность:", self.OOTIKZI_POSITION),
                ("initials", "Краткое имя:", self.OOTIKZI_INITIALS),
            ]
        ):
            ttk.Label(fo, text=label, anchor="w", width=16).grid(
                row=i, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=default)
            self._ootikzi[key] = var
            ttk.Entry(fo, textvariable=var, width=52).grid(
                row=i, column=1, padx=4, sticky="ew"
            )

        # ── Данные пользователя ───────────────────────
        fu = ttk.LabelFrame(main, text="  Данные пользователя  ", padding=8)
        fu.pack(fill="x", **PAD)
        fu.columnconfigure(1, weight=1)
        self._usr = {}
        for i, (key, label, default) in enumerate(
            [
                ("full_name", "ФИО (именительный):", ""),
                ("position", "Должность:", ""),
                ("genitive", "ФИО (родительный):", ""),
                ("initials", "Краткое имя (подписи):", ""),
                ("dept_abbrev", "Аббр. отдела (подписи):", ""),
                ("sticker", "№ пломбы-наклейки:", ""),
            ]
        ):
            ttk.Label(fu, text=label, anchor="w", width=26).grid(
                row=i, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=default)
            self._usr[key] = var
            ttk.Entry(fu, textvariable=var, width=46).grid(
                row=i, column=1, padx=4, sticky="ew"
            )

        # ── Реквизиты документа ───────────────────────
        fd = ttk.LabelFrame(main, text="  Реквизиты Акта  ", padding=8)
        fd.pack(fill="x", **PAD)
        fd.columnconfigure(1, weight=1)
        self._doc = {}
        now = datetime.now()
        for i, (key, label, default) in enumerate(
            [
                ("reg_number", "Рег. номер Акта:", ""),
                ("date_short", "Дата (дд.мм.гггг):", now.strftime("%d.%m.%Y")),
                (
                    "skzi_name",
                    "Название СКЗИ:",
                    "СКЗИ «ViPNet CSP» в составе ПО «ViPNet PKI Client»",
                ),
                ("skzi_version", "Версия СКЗИ:", "1.6"),
                ("skzi_build", "Сборка СКЗИ:", "1.542"),
                ("skzi_inventory", "Инвентарный № СКЗИ:", "14/852-ОСФР"),
                ("executor", "Исполнитель (журнал):", ""),
                ("output_dir", "Папка сохранения:", os.path.expanduser("~/Desktop")),
            ]
        ):
            ttk.Label(fd, text=label, anchor="w", width=26).grid(
                row=i, column=0, sticky="w", pady=2
            )
            var = tk.StringVar(value=default)
            self._doc[key] = var
            ttk.Entry(fd, textvariable=var, width=44).grid(
                row=i, column=1, padx=4, sticky="ew"
            )
            if key == "output_dir":
                ttk.Button(fd, text="…", width=3, command=self._pick_outdir).grid(
                    row=i, column=2
                )

        ttk.Button(
            main, text="📄   Создать Акт и добавить в журнал", command=self._generate
        ).pack(pady=10, ipadx=16, ipady=6)

        self._load_status_var = tk.StringVar(
            value="Укажите файлы данных и нажмите «Загрузить»"
        )
        self._load_status_lbl = ttk.Label(
            main, textvariable=self._load_status_var, foreground="#555", wraplength=700
        )
        self._load_status_lbl.pack(pady=(0, 12))

    # ── Helpers ───────────────────────────────────────

    def _set_status(self, text, color="gray"):
        clr = {
            "green": "#1a7a1a",
            "red": "#cc0000",
            "orange": "#b06000",
            "gray": "#555",
        }
        self.status_var.set(text)
        self._status_lbl.configure(foreground=clr.get(color, "#555"))

    def _set_load_status(self, text, color="gray"):
        clr = {"green": "#2a7a2a", "red": "#cc0000", "gray": "#555"}
        self._load_status_var.set(text)
        self._load_status_lbl.configure(foreground=clr.get(color, "#555"))

    def _pick_file(self, key, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self._file_vars[key].set(path)
            self.cfg[key] = path
            save_config(self.cfg)

    def _pick_outdir(self):
        d = filedialog.askdirectory()
        if d:
            self._doc["output_dir"].set(d)

    def _open_file(self, path):
        try:
            if sys.platform == "darwin":
                subprocess.run(["open", path])
            elif sys.platform == "win32":
                os.startfile(path)
            else:
                subprocess.run(["xdg-open", path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть: {e}")

    # ── Load ──────────────────────────────────────────

    def _load_files(self):
        paths = {k: self._file_vars[k].get().strip() for k in self._file_vars}
        errors = []
        loaded = 0

        if paths["phone_book"] and os.path.exists(paths["phone_book"]):
            try:
                self.phone_data = load_phone_book(paths["phone_book"])
                loaded += 1
            except Exception as e:
                errors.append(f"Справочник: {e}")

        if paths["pc_file"] and os.path.exists(paths["pc_file"]):
            try:
                self.pc_data = load_pc_data(paths["pc_file"])
                loaded += 1
            except Exception as e:
                errors.append(f"Актуализация ПК: {e}")

        # Объединяем имена из обоих источников
        self.all_names = sorted(set(self.phone_data.keys()) | set(self.pc_data.keys()))

        if paths["journal"] and os.path.exists(paths["journal"]):
            try:
                self.journal_info = get_journal_info(paths["journal"])
                self._doc["reg_number"].set(self.journal_info["next_reg"])
                if self.journal_info["last_executor"]:
                    self._doc["executor"].set(self.journal_info["last_executor"])
                loaded += 1
            except Exception as e:
                errors.append(f"Журнал: {e}")

        for k in paths:
            self.cfg[k] = paths[k]
        save_config(self.cfg)

        if errors:
            self._set_load_status("⚠ " + " | ".join(errors), "red")
        else:
            ji = self.journal_info
            self._set_load_status(
                f"✓ Загружено {loaded}/3 файлов. "
                f"Сотрудников: {len(self.all_names)}. "
                f"Следующий рег. номер: {ji.get('next_reg','—')}",
                "green",
            )

    # ── Search employee ───────────────────────────────

    def _on_search_change(self, *_):
        q = self.search_var.get().strip().lower()
        self.lb.delete(0, tk.END)
        if len(q) < 2:
            return
        for name in self.all_names:
            if q in name.lower():
                self.lb.insert(tk.END, name)

    def _on_select(self, _):
        sel = self.lb.curselection()
        if not sel:
            return
        name = self.lb.get(sel[0])
        self.search_var.set(name)
        self._fill_employee(name)

    def _fill_employee(self, name: str):
        emp = self.phone_data.get(name, {})
        position = emp.get("position", "")
        dept = emp.get("department", "")
        abbrev = abbreviate_dept(dept)

        self._usr["full_name"].set(name)
        self._usr["position"].set(build_position_doc(position, dept))
        self._usr["initials"].set(make_initials(name))
        self._usr["dept_abbrev"].set(abbrev)
        self._usr["genitive"].set(_to_genitive(name))

        # Если сотрудник найден только в актуализации — предупреждаем
        if not emp:
            messagebox.showwarning(
                "Сотрудник не в телефонном справочнике",
                f"\u00ab{name}\u00bb найден только в таблице актуализации.\n\n"
                "Заполните вручную следующие поля:\n"
                "  \u2022 Должность и отдел\n"
                "  \u2022 Начальник (ФИО)\n"
                "  \u2022 Инициалы начальника\n"
                "  \u2022 Аббревиатура отдела",
            )

        canonical = _strip_suffix(name)
        pcs = self.pc_data.get(canonical, []) or self.pc_data.get(name, [])
        self._pcs = pcs

        if not pcs:
            self.pc_combo["values"] = ["Не найден в таблице актуализации"]
            self.pc_combo.current(0)
            self._serial_var.set("")
            self._inv_var.set("")
            self._floor_var.set("")
            self._room_var.set("")
            self._usr["sticker"].set("")
            self._preview.clear()
            self._set_status("⚠ ПК не найден в таблице актуализации", "orange")
        else:
            labels = [pc["label"] for pc in pcs]
            self.pc_combo["values"] = labels
            self.pc_combo.current(0)
            self._fill_pc(pcs[0])
            self._preview.show(pcs)
            if len(pcs) == 1:
                self._set_status(f"✓ Найден: {name}", "green")
            else:
                self._set_status(
                    f"⚠ {len(pcs)} устройства — выберите в «Выбор ПК»", "orange"
                )

    def _fill_pc(self, pc: dict):
        self._serial_var.set(pc["serial"])
        self._inv_var.set(pc["inventory"])
        vals = pc.get("values", [])
        self._floor_var.set(vals[0].strip() if len(vals) > 0 else "")
        self._room_var.set(vals[1].strip() if len(vals) > 1 else "")
        self._usr["sticker"].set(vals[10].strip() if len(vals) > 10 else "")

    def _on_pc_select(self, _):
        idx = self.pc_combo.current()
        if 0 <= idx < len(self._pcs):
            self._fill_pc(self._pcs[idx])

    def _on_journal_search(self):
        journal_path = self._file_vars["journal"].get().strip()
        if not journal_path or not os.path.exists(journal_path):
            self._jstatus_var.set("Журнал не указан или не найден")
            return
        query = self._jsearch_var.get().strip()
        if len(query) < 2:
            for row in self._jtree.get_children():
                self._jtree.delete(row)
            self._jstatus_var.set("")
            return
        # Кеш: перечитываем журнал только если путь изменился
        if (
            not hasattr(self, "_journal_cache")
            or self._journal_cache_path != journal_path
        ):
            self._journal_cache = search_journal_entries(
                journal_path, ""
            )  # загружаем всё
            self._journal_cache_path = journal_path
        q = query.lower()
        results = [
            r
            for r in self._journal_cache
            if q in r.get("desc", "").lower() or q in r.get("executor", "").lower()
        ]
        # Очистить таблицу
        for row in self._jtree.get_children():
            self._jtree.delete(row)
        if not results:
            self._jstatus_var.set(f"Ничего не найдено по «{query}»")
        else:
            for r in results:
                self._jtree.insert(
                    "",
                    "end",
                    values=(r["pp"], r["date"], r["reg"], r["desc"], r["executor"]),
                )
            self._jstatus_var.set(
                f"Найдено {len(results)} запись(-ей) — кликните на нужную"
            )

    def _on_jresult_select(self, _):
        sel = self._jtree.selection()
        if not sel:
            return
        vals = self._jtree.item(sel[0], "values")
        if len(vals) >= 3:
            reg = vals[2]
            date = vals[1]
            ref = f"заявки {reg} от {date} г" if reg else vals[3][:60]
            self._zayavka_ref_var.set(ref)

    # ── Generate ──────────────────────────────────────

    def _ask_duplicate(self, path: str, filename: str) -> str:
        win = tk.Toplevel(self)
        win.title("Файл уже существует")
        win.resizable(False, False)
        win.grab_set()
        ttk.Label(
            win,
            text=f"Акт уже существует:\n{filename}",
            wraplength=400,
            justify="left",
            padding=12,
        ).pack()
        result = tk.StringVar(value="cancel")
        bf = ttk.Frame(win, padding=8)
        bf.pack()

        def choose(v):
            result.set(v)
            win.destroy()

        ttk.Button(
            bf, text="📂  Открыть существующий", command=lambda: choose("open")
        ).grid(row=0, column=0, padx=6, pady=4, sticky="ew")
        ttk.Button(bf, text="🔄  Заменить", command=lambda: choose("replace")).grid(
            row=0, column=1, padx=6, pady=4, sticky="ew"
        )
        ttk.Button(bf, text="➕  Создать ещё один", command=lambda: choose("new")).grid(
            row=1, column=0, padx=6, pady=4, sticky="ew"
        )
        ttk.Button(bf, text="✖  Отмена", command=lambda: choose("cancel")).grid(
            row=1, column=1, padx=6, pady=4, sticky="ew"
        )
        win.wait_window()
        return result.get()

    def _generate(self):
        user_name = self._usr["full_name"].get().strip()
        if not user_name:
            messagebox.showerror("Ошибка", "Сотрудник не выбран!")
            return

        reg = self._doc["reg_number"].get().strip()
        if not reg:
            messagebox.showerror("Ошибка", "Укажите регистрационный номер!")
            return

        date_str = self._doc["date_short"].get().strip()
        try:
            dt = datetime.strptime(date_str, "%d.%m.%Y")
        except ValueError:
            messagebox.showerror("Ошибка", "Формат даты: дд.мм.гггг")
            return

        genitive = self._usr["genitive"].get().strip()
        if not genitive:
            messagebox.showerror("Ошибка", "Укажите ФИО в родительном падеже!")
            return

        serial = self._serial_var.get().strip()
        inventory = self._inv_var.get().strip()
        if not serial and not inventory:
            if not messagebox.askyesno(
                "Предупреждение",
                "Серийный и инвентарный номера не заполнены.\nПродолжить?",
            ):
                return

        emp_initials = make_initials(user_name)

        data = {
            "reg_number": reg,
            "day": str(dt.day),
            "month": MONTHS_RU[dt.month],
            "year": str(dt.year),
            "zayavka_ref": self._zayavka_ref_var.get().strip(),
            "floor": self._floor_var.get().strip(),
            "room": self._room_var.get().strip(),
            "serial": serial,
            "user_full_name": user_name,
            "user_full_name_genitive": genitive,
            "user_position": self._usr["position"].get().strip(),
            "user_initials": self._usr["initials"].get().strip(),
            "user_dept_abbrev": self._usr["dept_abbrev"].get().strip(),
            "ootikzi_full_name": self._ootikzi["name"].get().strip(),
            "ootikzi_position": self._ootikzi["position"].get().strip(),
            "ootikzi_initials": self._ootikzi["initials"].get().strip(),
            "sticker_number": self._usr["sticker"].get().strip(),
            "skzi_name": self._doc["skzi_name"].get().strip()
            or "СКЗИ «ViPNet CSP» в составе ПО «ViPNet PKI Client»",
            "skzi_version": self._doc["skzi_version"].get().strip() or "1.6",
            "skzi_build": self._doc["skzi_build"].get().strip() or "1.542",
            "skzi_inventory": self._doc["skzi_inventory"].get().strip()
            or "14/852-ОСФР",
        }

        out_dir = self._doc["output_dir"].get().strip() or os.path.expanduser(
            "~/Desktop"
        )
        os.makedirs(out_dir, exist_ok=True)
        base_name = (
            f"Акт установки СКЗИ ViPNet CSP в составе ПО ViPNet PKI Client "
            f"{emp_initials} ({serial}).docx"
        )
        journal_desc_default = (
            f"Акт установки СКЗИ ViPNet CSP в составе "
            f"ПО ViPNet PKI Client {emp_initials} ({serial})"
        )

        # ── Окно подтверждения ────────────────────────
        confirm = ConfirmWindow(self, base_name, journal_desc_default, out_dir)
        confirmed_name, confirmed_desc = confirm.result
        if confirmed_name is None:
            return  # пользователь нажал Отмена
        base_name = confirmed_name
        output_path = os.path.join(out_dir, base_name)

        if os.path.exists(output_path):
            choice = self._ask_duplicate(output_path, base_name)
            if choice == "open":
                self._open_file(output_path)
                return
            elif choice == "cancel":
                return
            elif choice == "new":
                ts = datetime.now().strftime("%H-%M-%S")
                output_path = os.path.join(
                    out_dir, base_name.replace(".docx", f"_{ts}.docx")
                )

        # Хуки для подклассов (переопределяются в AktCSPApp и др.)
        data.setdefault("p2_suffix", "")
        if hasattr(self, "_override_data_hooks"):
            data = self._override_data_hooks(data)
        try:
            generate_akt_pki(data, output_path)
        except Exception as e:
            messagebox.showerror("Ошибка при создании Акта", str(e))
            return

        journal_path = self._file_vars["journal"].get().strip()
        journal_msg = ""
        if journal_path and os.path.exists(journal_path) and self.journal_info:
            executor_name = self._doc["executor"].get().strip()
            acquired, blocker = acquire_journal_lock(journal_path, executor_name)
            if not acquired:
                if blocker is None:
                    messagebox.showwarning(
                        "Ошибка записи в журнал",
                        "Не удалось заблокировать журнал для записи.\n"
                        "Папка с журналом, возможно, открыта только для чтения.\n"
                        "Обратитесь к администратору.",
                    )
                    journal_msg = (
                        "⚠ Нет прав на запись рядом с журналом — запись не добавлена."
                    )
                else:
                    messagebox.showwarning(
                        "Журнал заблокирован",
                        f"Журнал сейчас редактирует другой пользователь:\n{blocker}\n\n"
                        "Подождите и попробуйте снова.",
                    )
                    journal_msg = (
                        f"⚠ Журнал заблокирован: {blocker} — запись не добавлена."
                    )
            else:
                try:
                    entry = {
                        "pp": self.journal_info["next_pp"],
                        "date": date_str,
                        "reg": reg,
                        "description": confirmed_desc,
                        "executor": executor_name,
                        "note": "",
                    }
                    journal_msg = _write_journal_with_retry(
                        journal_path, entry, self.journal_info["last_row_idx"]
                    )
                    if journal_msg.startswith("✓"):
                        self.journal_info = get_journal_info(journal_path)
                        self._doc["reg_number"].set(self.journal_info["next_reg"])
                finally:
                    release_journal_lock(journal_path)
        else:
            journal_msg = "Журнал не указан — запись не добавлена."

        self._set_status(f"✓ Создан: {os.path.basename(output_path)}", "green")
        if messagebox.askyesno(
            "Готово!", f"Акт создан:\n{output_path}\n\n{journal_msg}\n\nОткрыть файл?"
        ):
            self._open_file(output_path)


# ══════════════════════════════════════════════════════
#  AktCSPApp — Акт установки СКЗИ ViPNet CSP Client
# ══════════════════════════════════════════════════════


class AktCSPApp(AktPKIApp):
    """Акт установки ViPNet CSP Client — наследует AktPKIApp, отличия только в дефолтах."""

    CSP_SKZI = "СКЗИ ViPNet CSP в составе СПО ViPNet Client c Деловой почтой"
    CSP_VER = "4.5"
    CSP_BUILD = "3.65160"
    CSP_INV = "1044-ОСФР"

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Акт установки СКЗИ ViPNet CSP Client")
        # Переопределяем дефолты полей после инициализации родителя
        self._doc["skzi_name"].set(self.CSP_SKZI)
        self._doc["skzi_version"].set(self.CSP_VER)
        self._doc["skzi_build"].set(self.CSP_BUILD)
        self._doc["skzi_inventory"].set(self.CSP_INV)

    def _override_data_hooks(self, data):
        data["p2_suffix"] = " и проверка получения обновлений из УЦ ОПФР"
        data["p3_genitive_prefix"] = "у "
        data["pak_center"] = True
        return data


# ══════════════════════════════════════════════════════
#  AktCSPOnlyApp — Акт установки СКЗИ ViPNet CSP (без Client)
# ══════════════════════════════════════════════════════


class AktCSPOnlyApp(AktPKIApp):
    """Акт установки ViPNet CSP — без 'в составе СПО...', без УЦ ОПФР, с 'у ФИО'."""

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Акт установки СКЗИ ViPNet CSP")
        self._doc["skzi_name"].set("СКЗИ ViPNet CSP")
        self._doc["skzi_version"].set("4.4")
        self._doc["skzi_build"].set("8.7899")
        self._doc["skzi_inventory"].set("908-ОСФР")

    def _override_data_hooks(self, data):
        data["p3_genitive_prefix"] = "у "
        data["pak_center"] = True
        return data


# ══════════════════════════════════════════════════════
#  AktKriptoproApp — Акт установки СКЗИ КриптоПро CSP
# ══════════════════════════════════════════════════════


class AktKriptoproApp(AktPKIApp):
    """Акт установки КриптоПро CSP."""

    def __init__(self, launcher):
        super().__init__(launcher)
        self.title("Акт установки СКЗИ КриптоПро CSP")
        self._doc["skzi_name"].set("«СКЗИ КриптоПро CSP»")
        self._doc["skzi_version"].set("5.0")
        self._doc["skzi_build"].set("1300")
        self._doc["skzi_inventory"].set("1045-ОСФР")

    def _override_data_hooks(self, data):
        data["p3_genitive_prefix"] = "у "
        data["pak_center"] = True
        return data


# ══════════════════════════════════════════════════════
#  LauncherWindow — главный экран выбора
# ══════════════════════════════════════════════════════


class LauncherWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Генератор заявок ViPNet")
        self.geometry("600x750")
        self.minsize(560, 500)
        self.resizable(True, True)
        if sys.platform == "darwin":
            self._setup_mac_clipboard()
        elif sys.platform == "win32":
            self._setup_win_clipboard()
        self._build()

    def _setup_mac_clipboard(self):
        root = self

        def _copy(e):
            try:
                sel = e.widget.selection_get()
                root.clipboard_clear()
                root.clipboard_append(sel)
            except Exception:
                pass
            return "break"

        def _paste(e):
            try:
                text = root.clipboard_get()
                try:
                    e.widget.delete("sel.first", "sel.last")
                except Exception:
                    pass
                e.widget.insert("insert", text)
            except Exception:
                pass
            return "break"

        def _cut(e):
            try:
                sel = e.widget.selection_get()
                root.clipboard_clear()
                root.clipboard_append(sel)
                e.widget.delete("sel.first", "sel.last")
            except Exception:
                pass
            return "break"

        def _select_all(e):
            try:
                e.widget.select_range(0, "end")
                e.widget.icursor("end")
            except Exception:
                pass
            return "break"

        self.bind_class("Entry", "<Command-c>", _copy)
        self.bind_class("Entry", "<Command-v>", _paste)
        self.bind_class("Entry", "<Command-x>", _cut)
        self.bind_class("Entry", "<Command-a>", _select_all)

    def _setup_win_clipboard(self):
        root = self

        def _copy(e):
            try:
                sel = e.widget.selection_get()
                root.clipboard_clear()
                root.clipboard_append(sel)
            except Exception:
                pass
            return "break"

        def _paste(e):
            try:
                text = root.clipboard_get()
                try:
                    e.widget.delete("sel.first", "sel.last")
                except Exception:
                    pass
                e.widget.insert("insert", text)
            except Exception:
                pass
            return "break"

        def _cut(e):
            try:
                sel = e.widget.selection_get()
                root.clipboard_clear()
                root.clipboard_append(sel)
                e.widget.delete("sel.first", "sel.last")
            except Exception:
                pass
            return "break"

        def _select_all(e):
            try:
                e.widget.select_range(0, "end")
                e.widget.icursor("end")
            except Exception:
                pass
            return "break"

        self.bind_class("Entry", "<Control-c>", _copy)
        self.bind_class("Entry", "<Control-v>", _paste)
        self.bind_class("Entry", "<Control-x>", _cut)
        self.bind_class("Entry", "<Control-a>", _select_all)

    def _build(self):
        canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0, yscrollincrement=20)
        vbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        inner = tk.Frame(canvas)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")
        inner.bind("<Configure>", lambda _: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(win_id, width=e.width))

        def _on_mousewheel(event):
            if not canvas.winfo_exists():
                return
            # Windows передаёт delta кратно 120, иногда меньше
            if event.delta:
                d = int(-1 * (event.delta / 120))
                if d == 0:
                    d = -1 if event.delta > 0 else 1
                canvas.yview_scroll(d, "units")
        self.bind_all("<MouseWheel>", _on_mousewheel)
        self.bind_all("<Button-4>", lambda e: canvas.winfo_exists() and canvas.yview_scroll(-1, "units"))
        self.bind_all("<Button-5>", lambda e: canvas.winfo_exists() and canvas.yview_scroll(1, "units"))

        ttk.Label(
            inner, text="Выберите тип документа", font=("TkDefaultFont", 14, "bold")
        ).pack(pady=(16, 4))

        def _card(parent, row, col, title, subtitle, cmd):
            fr = tk.Frame(
                parent, relief="groove", bd=2, padx=12, pady=12, cursor="hand2"
            )
            fr.grid(row=row, column=col, padx=8, pady=6, sticky="nsew")
            tk.Label(
                fr,
                text=title,
                font=("TkDefaultFont", 11, "bold"),
                wraplength=220,
                justify="center",
            ).pack()
            tk.Label(
                fr,
                text=subtitle,
                font=("TkDefaultFont", 9),
                foreground="#555",
                wraplength=220,
                justify="center",
            ).pack(pady=(5, 0))
            ttk.Button(fr, text="Открыть →", command=cmd).pack(pady=(10, 0))

        # ── Секция заявок ─────────────────────────────
        ttk.Label(
            inner,
            text="Заявки на установку СКЗИ",
            font=("TkDefaultFont", 10, "bold"),
            foreground="#555",
        ).pack(anchor="w", padx=24)

        cards = tk.Frame(inner)
        cards.pack(padx=16, fill="x")
        cards.columnconfigure(0, weight=1)
        cards.columnconfigure(1, weight=1)

        _card(
            cards,
            0,
            0,
            "ViPNet PKI Client",
            "Заявка на обучение и установку PKI Client",
            lambda: self._open_app(PKIApp),
        )
        _card(
            cards,
            0,
            1,
            "ViPNet CSP Client",
            "Заявка на установку СКЗИ ViPNet CSP / ViPNet Client",
            lambda: self._open_app(CSPApp),
        )
        _card(
            cards,
            1,
            0,
            "ViPNet CSP (ПТК КС)",
            "Заявка на установку СКЗИ ViPNet CSP для работы в ПТК КС",
            lambda: self._open_app(CSPPtkApp),
        )
        _card(
            cards,
            1,
            1,
            "КриптоПРО CSP",
            "Заявка на обучение СКЗИ КриптоПРО CSP",
            lambda: self._open_app(KriptoproApp),
        )

        # ── Секция актов ──────────────────────────────
        ttk.Separator(inner, orient="horizontal").pack(fill="x", padx=16, pady=(8, 4))
        ttk.Label(
            inner,
            text="Акты установки СКЗИ",
            font=("TkDefaultFont", 10, "bold"),
            foreground="#555",
        ).pack(anchor="w", padx=24)

        acts = tk.Frame(inner)
        acts.pack(padx=16, fill="x")
        acts.columnconfigure(0, weight=1)
        acts.columnconfigure(1, weight=1)

        _card(
            acts,
            0,
            0,
            "Акт установки PKI",
            "Акт установки СКЗИ ViPNet CSP в составе ПО ViPNet PKI Client",
            lambda: self._open_app(AktPKIApp),
        )
        _card(
            acts,
            0,
            1,
            "Акт установки CSP Client",
            "Акт установки СКЗИ ViPNet CSP в составе СПО ViPNet Client c Деловой почтой",
            lambda: self._open_app(AktCSPApp),
        )
        _card(
            acts,
            1,
            0,
            "Акт установки CSP",
            "Акт установки СКЗИ ViPNet CSP",
            lambda: self._open_app(AktCSPOnlyApp),
        )
        _card(
            acts,
            1,
            1,
            "Акт установки КриптоПро",
            "Акт установки СКЗИ КриптоПро CSP",
            lambda: self._open_app(AktKriptoproApp),
        )

        # ── Секция ЕЦП ────────────────────────────────
        ttk.Separator(inner, orient="horizontal").pack(fill="x", padx=16, pady=(8, 4))
        ttk.Label(
            inner,
            text="ЕЦП",
            font=("TkDefaultFont", 10, "bold"),
            foreground="#555",
        ).pack(anchor="w", padx=24)

        ecp_frame = tk.Frame(inner)
        ecp_frame.pack(padx=16, fill="x")
        ecp_frame.columnconfigure(0, weight=1)
        ecp_frame.columnconfigure(1, weight=1)

        _card(
            ecp_frame,
            0,
            0,
            "Заявка ЕЦП",
            "Заявка на предоставление доступа в ГИС ЕЦП",
            lambda: self._open_app(ECPApp),
        )

    def _open_app(self, AppClass):
        self.withdraw()
        AppClass(self)


if __name__ == "__main__":
    launcher = LauncherWindow()
    launcher.mainloop()